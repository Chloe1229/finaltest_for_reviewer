import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from copy import deepcopy
from tempfile import NamedTemporaryFile
import os
import textwrap
import re
BASE_DIR = os.path.dirname(__file__)


def render_step_title(num):
    st.markdown(
        f"<h2 style='font-size:20px; white-space:nowrap; text-align:center; background-color:#FAF3DB; padding:4px;'>「의약품 허가 후 제조방법 변경관리 가이드라인」 적용 대상 확인 {num}</h2>",
        unsafe_allow_html=True,
    )
    # add a blank line below the title for spacing
    st.markdown("<br>", unsafe_allow_html=True)


def render_cover_page():
    """Render the introductory cover page before step1."""
    # initial spacing from the very top
    st.markdown("<br>", unsafe_allow_html=True)

    cover_html = """
<div style="background-color:#FAF3DB; padding:20px; text-align:center;">
  <p>───────────────────────────────────────</p>
  <p style="font-size:14pt; margin:0;">의약품 제조방법 변경관리 지원도구 (비공식)<br><br></p>
  <p style="font-size:17pt; margin:0;">「의약품 허가 후 제조방법 변경관리 가이드라인(2024.12)" 기반의<br>
      보고유형 및 필요서류 자동 분류 및 제출양식 생성 시스템</p>
  <p>───────────────────────────────────────</p>
</div>
    """
    st.markdown(cover_html, unsafe_allow_html=True)

    # add space below the title box
    st.markdown("<br>", unsafe_allow_html=True)

    # start button centered below the title box
    col_left, col_button, col_right = st.columns([2, 1, 2])
    with col_button:
        if st.button("▶️시작하기"):
            st.session_state.step = 1

    # spacing after the start button
    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        """
<p style='text-indent:-1.4em; margin-left:1.4em;'>▫️ 이 시스템은 식품의약품안전처에서 공표한 자료를 바탕으로 개발된 개인 연구용 도구이며, 규제기관의 공식 입장이나 판단을 대변하지 않습니다.</p>
<p style='text-indent:-1.4em; margin-left:1.4em;'>▫️ 일부 기술적 표현(예: ‘∼하여야 한다’)이 포함되어 있더라도, 본 도구는 법적 효력을 가지지 않으며, 반드시 준수하여야 하는 사항으로 간주되어서는 안 됩니다.</p>
<p style='text-indent:-1.4em; margin-left:1.4em;'>▫️ 또한, 본 시스템의 모든 내용은 2025년 6월 21일 현재 유효한 가이드라인을 기준으로 작성되었으나, 이후 개정되는 법령, 고시, 또는 가이드라인에 따라 적용 내용이 달라질 수 있습니다.</p>
<p style='text-indent:-1.4em; margin-left:1.4em;'>▫️ 사용자는 본 도구를 참고자료로 활용하시되, 실제 규제 제출 또는 심사에 앞서 반드시 최신 법령 및 관련 규정을 별도로 확인하시기 바랍니다.</p>
<p style='text-indent:-1.4em; margin-left:1.4em;'>📚 출처: 「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」, 2024.12 / 일부 재구성<br>적용 기준일: <strong>2025년 6월 21일 현재 유효한 가이드라인 기준</strong></p>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)

    # copyright notice moved to the bottom
    st.markdown(
        "<div style='text-align:center; font-size:10pt;'>ⓒ 2025 Chloe Kim. All rights reserved.<br>본 시스템의 자동화 로직, UI 구성, 데이터 분류 구조는 Chloe Kim의 창작물로 보호되며,  무단 복제 및 2차 사용을 금합니다.</div>",
        unsafe_allow_html=True,
    )

def render_footer():
    """Render the copyright notice at the bottom of pages."""
    st.markdown(
        "<div style='text-align:center; font-size:10pt;'>ⓒ 2025 Chloe Kim. All rights reserved.<br>본 시스템의 자동화 로직, UI 구성, 데이터 분류 구조는 Chloe Kim의 창작물로 보호되며,  무단 복제 및 2차 사용을 금합니다.</div>",
        unsafe_allow_html=True,
    )


# ===== 초기 상태 정의 =====
if "step" not in st.session_state:
    st.session_state.step = 0

if "step1_answer" not in st.session_state:
    st.session_state.step1_answer = None
if "step2_answer" not in st.session_state:
    st.session_state.step2_answer = None
if "step3_answer" not in st.session_state:
    st.session_state.step3_answer = None

if "step4_selections" not in st.session_state:
    st.session_state.step4_selections = {}

if "step5_targets" not in st.session_state:
    st.session_state.step5_targets = []
if "step9_dl1" not in st.session_state:
    st.session_state.step9_dl1 = False
if "step9_dl2" not in st.session_state:
    st.session_state.step9_dl2 = False

if st.session_state.step == 0:
    render_cover_page()
    st.stop()

# ===== Step1 함수 및 화면 =====
def go_to_step2():
    if st.session_state.step1_answer == "예":
        st.session_state.step = 2

if st.session_state.step == 1:
    render_step_title(1)
    st.write("**제6조제1항에 따라 국제공통기술문서(CTD)로 작성하여 허가를 받거나 신고한 의약품의 제조원 또는 제조방법을 변경하는 경우에 해당한다.**")

    st.session_state.step1_answer = st.radio("➤ 답변을 선택하세요.", ["예", "아니오"], key="step1_radio")

    if st.session_state.step1_answer == "예":
        st.success(
            textwrap.dedent(
                """\
💡CTD 작성대상 완제의약품 해당합니다.  
    (근거 : 「의약품의 품목허가·신고·심사 규정」제6조(국제공통기술문서 작성) 제1항,
    제3조의2(의약품의 허가ㆍ신고의 변경 처리) 제6항)"""
            )
        )
        st.button("다음단계로", on_click=go_to_step2)

    elif st.session_state.step1_answer == "아니오":
        st.warning(
            textwrap.dedent(
                """⚠️CTD 작성대상 완제의약품 해당여부를 확인하고, 작성 대상에 해당하는 경우 먼저, CTD 제3부 품질평가 자료 중
    3.2.S.2, 3.2.S.3 및 3.2.P.2, 3.2.P.3, 3.2.P.4, 3.2.P.7를 제출하여 제조방법 자료로서 심사 받으시기 바랍니다.
    (근거 : 「의약품의 품목허가·신고·심사 규정」제6조(국제공통기술문서 작성) 제1항,
    제3조의2(의약품의 허가ㆍ신고의 변경 처리) 제6항)"""
            )
        )

# ===== Step2 함수 및 화면 =====
def go_to_step3():
    if st.session_state.step2_answer == "예":
        st.session_state.step = 3

if st.session_state.step == 2:
    render_step_title(2)
    st.write("**제조에 관한 항목 (CTD 제3부 품질평가 자료 중 3.2.S.2, 3.2.S.3 및 3.2.P.2, 3.2.P.3, 3.2.P.4, 3.2.P.7)을 변경 하는 경우에 해당한다.**")

    st.session_state.step2_answer = st.radio("➤ 답변을 선택하세요.", ["예", "아니오"], key="step2_radio")

    if st.session_state.step2_answer == "예":
        st.success(
            textwrap.dedent(
                """\
💡 「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」의 적용 대상 항목의 변경에 해당합니다.
    (근거 : 「의약품의 품목허가·신고·심사 규정」[별표 19])"""
            )
        )
        st.button("다음단계로", on_click=go_to_step3)

    elif st.session_state.step2_answer == "아니오":
        st.warning(
            textwrap.dedent(
                """\
⚠️ 제조에 관한 항목은 CTD 제3부 품질평가 자료 중
    3.2.S.2, 3.2.S.3 및 3.2.P.2, 3.2.P.3, 3.2.P.4, 3.2.P.7에 해당하며
    「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」는 해당 항목에 대한 변경에 대해 안내하고 있으므로,
    가이드라인 적용 대상에 해당하지 않습니다.
    (근거 : 「의약품의 품목허가·신고·심사 규정」[별표 19])"""
            )
        )

# ===== Step3 함수 및 화면 =====
def go_to_step4():
    if st.session_state.step3_answer == "예":
        st.session_state.step = 4

if st.session_state.step == 3:
    render_step_title(3)
    st.write("**품목의 허가(신고) 사항 중 제조방법에 해당하는 자료(CTD 제3부 품질평가 자료 중 3.2.S.2, 3.2.S.3 및 3.2.P.2, 3.2.P.3, 3.2.P.4, 3.2.P.7)를 국제공통기술문서(CTD)로서 제출하여 심사받은 ‘제조방법 CTD 적용(또는 전환)’ 품목에 해당한다.**")

    st.session_state.step3_answer = st.radio("➤ 답변을 선택하세요.", ["예", "아니오"], key="step3_radio")

    if st.session_state.step3_answer == "예":
        st.success(
            textwrap.dedent(
                """\
💡 「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」에 따라 변경수준을 확인할 수 있습니다.
    (근거 : 「의약품의 품목허가·신고·심사 규정」[별표 19])"""
            )
        )
        st.button("다음단계로", on_click=go_to_step4)

    elif st.session_state.step3_answer == "아니오":
        st.warning(
            textwrap.dedent(
                """\
⚠️ 먼저, CTD 제3부 품질평가 자료 중 3.2.S.2, 3.2.S.3 및 3.2.P.2, 3.2.P.3, 3.2.P.4, 3.2.P.7를 제출하여 제조방법 자료로서 심사 받으시기 바랍니다.
    (근거 : 「의약품의 품목허가·신고·심사 규정」[별표 19])"""
            )
        )

# Step 4 상태 초기화
if "step4_selections" not in st.session_state:
    st.session_state.step4_selections = {}
if "step5_targets" not in st.session_state:
    st.session_state.step5_targets = []

# Step 4 항목 (프롬프트 문구 그대로)
step4_items = {
    "s1": "3.2.S.1 일반정보",
    "s2": "3.2.S.2 제조",
    "p1": "3.2.P.1 완제의약품의 성상 및 조성",
    "p3": "3.2.P.3 제조",
    "p4": "3.2.P.4 첨가제의 관리",
    "p7": "3.2.P.7 용기-마개 시스템",
    "ds": "디자인스페이스(Design Space)"
}

# Step 3 → Step 4 이동 함수
def go_to_step4():
    if st.session_state.step3_answer == "예":
        st.session_state.step = 4

# Step 4 → Step 5 이동 함수
def go_to_step5():
    st.session_state.step5_targets = [
        code for code, val in st.session_state.step4_selections.items() if val == "변경 있음"
    ]
    # reset selections to avoid leftovers when revisiting Step5
    st.session_state.step5_selections = {}
    st.session_state.step5_page = 0    
    # clear downstream state so Step6 recalculates from fresh data
    st.session_state.step6_targets = []
    st.session_state.step6_selections = {}
    st.session_state.step6_page = 0
    st.session_state.step7_results = {}
    st.session_state.step7_page = 0
    st.session_state.pop("step8_page", None)
    st.session_state.step = 5

# Step 4 이전단계 복귀 함수
def go_back_to_step3():
    st.session_state.step = 3

# Step 4 실행
if st.session_state.step == 4:
    indent = "&nbsp;" * 2
    st.markdown(
        "<h2 style='font-size:24px; text-align:center; background-color:#FAF3DB; padding:4px;'>가이드라인에 명시된 변경사항의 CTD 항목 선택</h2>",
        unsafe_allow_html=True,
    )
    # Add a blank line below the page title for better spacing
    st.write("")
    st.markdown(
        f"<p style='font-size:16px;'>⚠️ 가이드라인에 명시된 CTD 항목은 대표적인 번호를 기재한 것으로, 제출시 CTD 작성 지침에<br>{indent} 맞게 작성하여야 합니다</p>",
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <style>
        div[data-testid="stRadio"] > label p {
            font-size:16px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stRadio"] label div[data-testid="stMarkdownContainer"] p {
            font-size:16px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stMarkdownContainer"] p {
            margin-top:0;
            margin-bottom:0;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<h4 style='font-size:20px;'>3.2.S 원료의약품</h4>", unsafe_allow_html=True)
    col_s1, col_s2 = st.columns(2)
    with col_s1:
        st.markdown("<p style='font-size:18px;'>3.2.S.1 일반정보</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["s1"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_s1",
            label_visibility="collapsed"
        )
    with col_s2:
        st.markdown("<p style='font-size:18px;'>3.2.S.2 제조</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["s2"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_s2",
            label_visibility="collapsed"
        )

    st.markdown("<h4 style='font-size:20px;'>3.2.P 완제의약품</h4>", unsafe_allow_html=True)
    col_p1, col_p3 = st.columns(2)
    with col_p1:
        st.markdown("<p style='font-size:18px;'>3.2.P.1 완제의약품의 성상 및 조성</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["p1"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_p1",
            label_visibility="collapsed"
        )
    with col_p3:
        st.markdown("<p style='font-size:18px;'>3.2.P.3 제조</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["p3"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_p3",
            label_visibility="collapsed"
        )

    col_p4, col_p7 = st.columns(2)
    with col_p4:
        st.markdown("<p style='font-size:18px;'>3.2.P.4 첨가제의 관리</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["p4"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_p4",
            label_visibility="collapsed"
        )
    with col_p7:
        st.markdown("<p style='font-size:18px;'>3.2.P.7 용기-마개 시스템</p>", unsafe_allow_html=True)
        st.session_state.step4_selections["p7"] = st.radio(
            "",
            ["변경 있음", "변경 없음"],
            key="step4_radio_p7",
            label_visibility="collapsed"
        )

    st.markdown("<h4 style='font-size:20px;'>디자인스페이스</h4>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:18px;'>디자인스페이스(Design Space)</p>", unsafe_allow_html=True)
    st.session_state.step4_selections["ds"] = st.radio(
        "",
        ["변경 있음", "변경 없음"],
        key="step4_radio_ds",
        label_visibility="collapsed"
    )

    # 모든 항목 선택 여부 확인
    all_selected = all(
        v in ["변경 있음", "변경 없음"]
        for v in st.session_state.step4_selections.values()
    )

    col1, col2 = st.columns(2)
    with col1:
        st.button("이전단계로", on_click=go_back_to_step3)
    with col2:
        st.button("다음단계로", on_click=go_to_step5, disabled=not all_selected)

# ===== Step5 상태 초기화 =====
if "step5_selections" not in st.session_state:
    st.session_state.step5_selections = {}
if "step6_targets" not in st.session_state:
    st.session_state.step6_targets = []
if "step5_page" not in st.session_state:
    st.session_state.step5_page = 0

# ===== Step5 항목 정의 (전체 하드코딩) =====
step5_items = {
    "s1": {
        "title": "3.2.S.1 일반정보",
        "items": {
            "1": "1. 원료의약품 명칭변경"
        }
    },
    "s2": {
        "title": "3.2.S.2 제조",
        "items": {
            "2": "2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가",
            "3": "3. 원료의약품 제조 공정의 변경",
            "4": "4. 원료의약품 제조 공정관리 규격의 변경",
            "5": "5. 원료의약품 또는 중간체의 제조 규모 변경",
            "6": "6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경"
        }
    },
    "p1": {
        "title": "3.2.P.1 완제의약품의 성상 및 조성",
        "items": {
            "7": "7. 완제의약품 중 고형 제제의 조성 변경",
            "8": "8. 완제의약품 중 고형 제제의 코팅층 무게 변경",
            "9": "9. 완제의약품 중 고형제제를 제외한 그 외 제형의 조성 변경",
            "10": "10. 완제의약품 중 고형제제를 제외한 그 외 제형에 쓰이는 착색제 또는 착향제의 종류와 분량의 변경"
        }
    },
    "p3": {
        "title": "3.2.P.3 제조",
        "items": {
            "11": "11. 정성적 또는 정량적인 조성과 평균 질량의 변경이 없는 성상의 변경(단 잉크, 그림, 글자체 등 식별표시를 위한 변경은 제외)",
            "12": "12. 완제의약품 제조공정 중 일부공정 제조소 또는 전체공정 제조소의, 추가 또는 변경",
            "13": "13. 비무균제제의 제조 규모 변경",
            "14": "14. 무균제제의 제조 규모 변경",
            "15": "15. 완제의약품의 제조공정 변경",
            "16": "16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경"
        }
    },
    "p4": {
        "title": "3.2.P.4 첨가제의 관리",
        "items": {
            "17": "17. 첨가제 기원의 변경",
            "18": "18. 별규에 해당하는 첨가제의 규격 또는 시험방법변경",
            "19": "19. 식약처장이 인정하는 공정서 규격으로 첨가제 규격의 변경"
        }
    },
    "p7": {
        "title": "3.2.P.7 용기-마개 시스템",
        "items": {
            "20": "20. 비무균제제의 직접용기 및 포장재질, 종류 변경",
            "21": "21. 무균제제의 직접용기 및 포장 재질, 종류 변경",
            "22": "22. 직접 포장의 규격 변경",
            "23": "23. 포장단위 변경"
        }
    },
    "ds": {
        "title": "디자인스페이스(Design Space)",
        "items": {
            "24": "24. 디자인스페이스(Design Space) 변경"
        }
    }
}

# ===== Step 간 이동 함수 =====
def go_to_step6():
    st.session_state.step6_targets = [
        key for key, val in st.session_state.step5_selections.items() if val == "변경 있음"
    ]
    # clear previous Step6 selections when recalculating targets
    st.session_state.step6_selections = {}
    st.session_state.step6_page = 0
    st.session_state.step7_results = {}
    st.session_state.step7_page = 0
    st.session_state.pop("step8_page", None)
    st.session_state.step = 6

def go_back_to_step4():
    st.session_state.step = 4

# ===== Step5 페이지 이동 함수 =====
def go_to_prev_step5_page():
    if st.session_state.step5_page > 0:
        st.session_state.step5_page -= 1

def go_to_next_step5_page():
    if st.session_state.step5_page < len(st.session_state.step5_targets) - 1:
        st.session_state.step5_page += 1

# ===== Step5 화면 =====
if st.session_state.step == 5:
    st.markdown(
        "<h2 style='font-size:25.5px; text-align:center; background-color:#FAF3DB; padding:4px;'>허가 후 제조방법 변경사항 선택</h2>",
        unsafe_allow_html=True,
    )
    st.markdown(
        """
        <style>
        div[data-testid="stRadio"] > label p {
            font-size:17px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stRadio"] label div[data-testid="stMarkdownContainer"] p {
            font-size:17px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stMarkdownContainer"] p {
            margin-top:0;
            margin-bottom:0;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    targets = st.session_state.step5_targets
    if not targets:
        st.warning("이전 단계에서 선택된 항목이 없습니다.")
    else:
        current_code = targets[st.session_state.step5_page]
        section = step5_items.get(current_code)

        if section:
            st.markdown(f"<h4 style='font-size:20px;'>{section['title']}</h4>", unsafe_allow_html=True)
            for num, label in section["items"].items():
                key = f"{current_code}_{num}"
                if current_code == "ds":
                    st.markdown(
                        f"<p style='font-size:18px;'>{label} → 변경 있음 (자동 선택됨)</p>",
                        unsafe_allow_html=True,
                    )
                    st.session_state.step5_selections[key] = "변경 있음"
                else:
                    radio_key = f"step5_radio_{key}"
                    if radio_key not in st.session_state:
                        st.session_state[radio_key] = None
                    st.markdown(f"<p style='font-size:18px;'>{label}</p>", unsafe_allow_html=True)
                    st.session_state.step5_selections[key] = st.radio(
                        "",
                        ["변경 있음", "변경 없음"],
                        key=radio_key,
                        label_visibility="collapsed",
                    )

    all_selected = all(
        v in ["변경 있음", "변경 없음"]
        for k, v in st.session_state.step5_selections.items()
        if not k.startswith("ds_")
    )

    col1, col2 = st.columns(2)
    with col1:
        st.button(
            "이전단계로",
            on_click=go_back_to_step4 if st.session_state.step5_page == 0 else go_to_prev_step5_page,
        )
    with col2:
        if st.session_state.step5_page == len(st.session_state.step5_targets) - 1:
            st.button("다음단계로", on_click=go_to_step6, disabled=not all_selected)
        else:
            st.button("다음항목 선택하기", on_click=go_to_next_step5_page)

if "step6_selections" not in st.session_state:
    st.session_state.step6_selections = {}

step6_items = {
    "s1_1": {
        "title": "3.2.S.1 일반정보\n1. 원료의약품 명칭변경",
        "subitems": {},
        "requirements": {
            "1": "1. 유효성분은 그대로 유지된다."
        }
    },
    "s2_2": {
        "title": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "subitems": {
            "2a": "2a. 원료의약품의 출발물질 생산",
            "2b": "2b. 원료의약품의 중간체 생산",
            "2c": "2c. 원료의약품의 생산"
        },
        "requirements": {
            "1": "1. 무균 원료의약품이 아니다.",
            "2": "2. 변경된 제조소의 원료의약품은 완제연계심사로서 DMF 품질심사를 완료하였다.",
            "3": "3. 원료의약품 규격에는 변경이 없다.",
            "4": "4. 출발 물질의 규격, 불순물 프로파일 및 합성경로는 변경이 없다.",
            "5": "5. 규격(공정 중 관리, 모든 원료의 분석 방법 포함), 제조 방법, 상세한 합성 경로 및 중간체 규격의 변경이 없다.",
            "6": "6. 원료의약품의 경우, 그 결정형이 동일하고, 입자크기가 중요한 경우 입자 크기 분포에 유의한 차이가 없다.",
            "7": "7. 원료의약품의 규격 (공정 중 관리, 모든 원료의 분석 방법 포함), 제조 방법 (배치 크기 포함) 및 상세한 합성 경로는 변경이 없다.",
            "8": "8. 사람 또는 동물 유래 물질이 사용되는 경우, BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위험에 대한 평가가 필요한 새로운 공급자를 이용하지 않는다.",
            "9": "9. 출발물질은 조품의 원료, 등록대상 원료의약품의 성분, 등록하고자 하는 원료의약품과 화학구조가 유사한 성분※(염류, 이성체 등)에 해당하지 않는다.\n※ 예시 : 세프메타졸(등록대상성분) → 세프메타졸에스터\n아토르바스타틴(등록대상성분) → 아토르바스타틴페네틸보로네이트\n로자탄(등록대상성분) → 트리틸로자탄",
            "10": "실제 제조장소의 변경이 없는 제조소의 명칭 변경이다(인수합병에 의한 상호변경 포함)."
        }
    },
    "s2_3": {
        "title": "3.2.S.2 제조\n3. 원료의약품 제조 공정의 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 원료의약품의 물리적 성질(결정형, 비결정형 등)에 변경이 없다.",
            "2": "2. 난용성 원료의약품의 경우, 결정다형이 동일하고 입자 크기가 중요한 경우 입자 크기 분포에 유의한 차이가 없다.",
            "3": "3. 사람 또는 동물 유래 물질이 사용되는 경우, 바이러스 안정성 평가 및 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위험에 대한 평가가 필요한 새로운 공정이 포함되지 않는다.",
            "4": "4. 합성경로의 변경이 없고(중간체는 동일하게 유지됨) 새로운 시약, 촉매 혹은 용매가 사용되지 않는다.",
            "5": "5. 원료의약품의 정성적 및 정량적 불순물 프로파일이나 물리 화학적 성질의 변경이 없다.",
            "6": "6. 무균 원료의약품의 멸균 또는 무균 공정에 영향을 미치지 않는다.",
            "7": "7. 최종 중간체 이전 공정의 변경이다.",
            "8": "8. 변경 전·후, 출발물질 규격, 중간체 규격 또는 원료의약품 규격의 변경이 없다.",
            "9": "9. 원료의약품 규격에 변경이 없다."
        }
    },
    "s2_4": {
        "title": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경",
        "subitems": {
            "4a": "4a. 공정관리 기준 강화",
            "4b": "4b. 공정관리 시험 및 기준 추가",
            "4c": "4c. 안전성이나 품질 문제로 인한 공정 관리 시험의 추가 또는 교체",
            "4d": "4d. 공정관리 시험 삭제",
            "4e": "4e. 공정관리 시험 기준 완화"
        },
        "requirements": {
            "1": "1. 해당 변경은 제조 중 예기치 않은 사례로 발생한 것이 아니다. (예: 안전성이 미확인된 새로운 불순물 확인 또는 총 불순물 기준의 변경 등)",
            "2": "2. 해당 변경은 현재 승인된 기준 범위 내에 있다.",
            "3": "3. 분석 절차는 동일하다. (분석절차의 경미한 변경은 허용)",
            "4": "4. 삭제되는 공정관리 시험은 품질에 영향을 미치지 않는다.",
            "5": "5. 해당 변경은 무균 원료의약품의 멸균 공정에 영향을 주지 않는다."
        }
    },
    "s2_5": {
        "title": "3.2.S.2 제조\n5. 원료의약품 또는 중간체의 제조 규모 변경",
        "subitems": {
            "5a": "5a. 제조 규모의 10배 이하 확대",
            "5b": "5b. 제조 규모의 축소",
            "5c": "5c. 제조 규모의 10배 초과 확대"
        },
        "requirements": {
            "1": "1. 제조방법 및 공정관리의 변경은 제조규모 변경에 따른 변경만을 포함한다.",
            "2": "2. 제조 공정의 재현성에 영향이 없다.",
            "3": "3. 제조과정에서 예상하지 못한 사유로 기준을 충족하지 못한 경우 또는 안정성 문제 때문에 발생한 변경이 아니다."
        }
    },
    "s2_6": {
        "title": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료의 규격변경",
        "subitems": {
            "6a": "6a. 규격 기준의 강화",
            "6b": "6b. 시험방법의 변경",
            "6c": "6c. 기준 및 시험방법 추가",
            "6d": "6d. 기준 또는 시험방법의 삭제",
            "6e": "6e. 안전성이나 품질관리 문제로 인한 기준의 추가 또는 교체",
            "6f": "6f. 원료약품(용매, 시약, 촉매 등)에 대한 기준 완화",
            "6g": "6g. 원료의약품 출발물질 및 핵심중간체(조품원료, 등록대상 원료의약품성분, 원료의약품과 화학구조가 유사한 성분 등)에 대한 기준 완화"
        },
        "requirements": {
            "1": "1. 제조과정에서 예상하지 못한 사유로 기준을 충족하지 못한 경우 또는 안정성 문제 때문에 발생한 변경이 아니다.",
            "2": "2. 모든 변경 사항은 현재의 허용 기준범위 내에 있다.",
            "3": "3. 시험방법의 변경은 없다.",
            "4": "4. 변경 후 시험방법은 동일한 분석 기술 또는 원리로 수행된다.",
            "5": "5. 변경 후 시험법 밸리데이션은 적절하게 수행되었으며, 이전 시험방법과 동등 이상이다.",
            "6": "6. 원료의약품의 총 불순물 기준 변경은 없으며, 새로 검출되는 불순물은 없다.",
            "7": "7. 변경은 유전독성 불순물의 관리전략에 영향을 미치지 않는다.",
            "8": "8. 변경되는 시험항목은 중요하지 않거나 승인된 대체 시험항목이 있다.",
            "9": "9. 회수용매의 기준과 관련되어 있지 않다."
        }
    },
    "p1_7": {
        "title": "3.2.P.1 완제의약품의 성상 및 조성\n7. 완제의약품 중 고형 제제의 조성 변경",
        "subitems": {
            "7a": "7a. 타르색소(황색4호 제외)의 종류 변경",
            "7b": "7b. 착색제 또는 착향제의 종류 및 분량 변경",
            "7c": "7c. 첨가제의 종류 및 분량 변경"
        },
        "requirements": {
            "1": "1. 제형의 기능적 특성 및 제품 품질(붕해시간 혹은 용출프로파일 등)에 미치는 영향은 없다.",
            "2": "2. 단위제형 총 중량 중 착색제, 착향제간의 함유율의 차는 없으며, 기허가 품목(동일투여경로)에 사용된 예가 있다 : 대한민국약전 또는 공정서에 수재된 성분 및 국내에서 사용례가 있는 성분과 이들 성분들로 조합된 혼합물질(착향제 포함), 식품첨가물의 기준 및 규격, 또는 일본의약품첨가물규격등 외국의 공인할 수 있는 자료 등에 의하여 사용 예를 인정할 수 있는 성분은 사용례를 인정 (의약품의 품목허가·신고·심사 규정 제25조제2항제1호)",
            "3": "3. 해당 변경은 안정성 문제로 인한 것이 아니며, 잠재적인 안전성 우려, 즉 용량 간의 구별 문제를 야기하지 않는다.",
            "4": "4. 변경/추가되는 첨가제는 「의약품의 품목허가·신고·심사 규정」 제25조제2항제1호에 따른 새로운 첨가제가 아니다.",
            "5": "5. 변경/추가되는 첨가제의 규격은 공정서 규격에 해당한다."
        }
    },
    "p1_8": {
        "title": "3.2.P.1 완제의약품의 성상 및 조성\n8. 고형제제의 코팅층 무게 변경",
        "subitems": {},
        "requirements": {}
    },
    "p1_9": {
        "title": "3.2.P.1 완제의약품의 성상 및 조성\n9. 완제의약품(고형제제 제외)의 조성 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 시럽제, 엘릭서제, 틴크제 등 경구용 액제 및 외용액제(유제 및 현탁제 제외)",
            "2": "2. 주사제, 점안제, 점이제로 원료약품의 종류가 이미 허가·신고사항과 동일하거나 다음의 첨가제가 다른 경우 : 「의약품의 품목허가·신고·심사 규정」 제27조제3항제2호에 따라 주사제는 보존제, 완충제, 항산화제, pH조절제(이미 허가·신고된 바있는 주사제에 사용된 pH조절제에 한함), 점안제 및 점이제는 보존제, 완충제, 등장화제, 점도조절제, pH조절제.",
            "3": "3. 흡입 전신마취제(「의약품의 품목허가·신고·심사 규정」 제27조제3항제3호)",
            "4": "4. 액제를 제외한 국소적용 외용제제로서 원료약품의 종류가 이미 허가·신고사항과 동일하거나 다음의 첨가제가 다른 경우(「의약품의 품목허가·신고·심사 규정」 제27조제3항제4호) : 보존제, 항산화제, 착색제 및 착향제",
            "5": "5. 유효성분을 기체나 증기 등의 흡입제로 투여하는 것으로서 국소요법만을 목적으로 하는 제제(「의약품의 품목허가·신고·심사 규정」 제27조제3항제5호)",
            "6": "6. 수액제, 혈액증량제 및 인공관류액제제(「의약품의 품목허가·신고·심사 규정」 제27조제3항제6호)",
            "7": "7. 폐에 적용하는 흡입제(「의약품의 품목허가·신고·심사 규정」 제27조제5항제1호단서조항)",
            "8": "8. 충족조건 1 ~ 7을 모두 충족하지 않는 경우",
            "9": "9. 변경/추가되는 첨가제는 「의약품의 품목허가·신고·심사 규정」 제25조제2항제1호에 따른 새로운 첨가제가 아니다."
        }
    },
    "p1_10": {
        "title": "3.2.P.1 완제의약품의 성상 및 조성\n10. 완제의약품(고형제제 제외)에 쓰이는 착색제 또는 착향제의 종류와 분량의 변경",
        "subitems": {
            "10a": "10a. 타르색소의 종류 변경(황색4호 제외)",
            "10b": "10b. 타르색소의 종류 및 분량 변경",
            "10c": "10c. 타르색소 외 착색제, 착향제의 변경"
        },
        "requirements": {
            "1": "1. 제형의 성능에 관한 시험결과(예: 붕해시간 및 용출프로파일)에 변화가 없다.",
            "2": "2. 총 중량을 유지하기 위해 주요 첨가제를 써서 처방 조성이 경미하게 조정되어 있다.",
            "3": "3. 완제의약품의 규격은 성상, 냄새 및/또는 맛에 관해서만 개정되어 있거나, 관련되는 경우, 확인시험이 삭제되거나 추가되어 있다.",
            "4": "4. 타르색소 착색제(황색 4호 제외)를 사용하는 경우 종류 및 분량은 「의약품등의 타르색소 지정과 기준 및 시험방법」(식약처 고시)에 적합하다.",
            "5": "5. 착향제는 기허가 품목(동일 투여경로)에 사용된 예가 있다.",
            "6": "6. 변경/추가되는 첨가제는 국제공통기술문서 3.2.P.4를 준수하고 있다.",
            "7": "7. 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위해 적합성 평가 자료가 있다.",
            "8": "8. 변경 사항은 용량 간 식별에 영향을 미치지 않는다.",
            "9": "9. 변경/추가되는 첨가제는 「의약품의 품목허가·신고·심사 규정」 제25조제2항제1호에 따른 새로운 첨가제가 아니다."
        }
    },
    "p3_11": {
        "title": "3.2.P.3 제조\n11. 정성적 또는 정량적인 조성과 평균 질량의 변경이 없는 성상의 변경(단, 식별표시를 위한 변경(잉크, 그림, 글자체 등) 제외)",
        "subitems": {
            "11a": "11a. 11b에 언급된 것 이외의 정제, 캡슐, 좌제",
            "11b": "11b. 장용성, 서방성 완제의약품"
        },
        "requirements": {
            "1": "완제의약품의 규격은 완제의약품의 모양에 대해서만 변경되어 있다.",
            "2": "변경 전 후 최소 1배치(파일럿 배치 이상)의 용출 프로파일(기준 및 시험방법 시험액에서 측정)은 동등하다.",
            "3": "3. 성형과정의 변경에 따른 모양 및/또는 크기의 변경이 있다."
        }
    },
    "p3_12": {
        "title": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경",
        "subitems": {
            "12a": "12a. 이차 포장 제조소",
            "12b1": "12b.1. 일차 포장 제조소 – 고형(정제, 캡슐제), 반고형(연고제, 크림제) 및 액상 제제",
            "12b2": "12b.2. 일차 포장 제조소 – 그 밖의 액상 제제(유제, 현탁제)",
            "12c1": "12c1. 원료칭량 공정 및 완제품 포장공정 제조소를 제외한 그 밖의 모든 제조소",
            "12c2": "12c2. 원료칭량 공정 및 완제품 포장공정 제조소를 제외한 그 밖의 모든 제조소"
        },
        "requirements": {
            "1": "1. 배치 조성, 제조 공정 및 공정 관리의 기록 사항, 장비 등급 및 공정 관리, 주요 공정 및 반제품의 관리 또는 완제의약품 규격에 있어서 변경이 없다.",
            "2": "2. (해당하는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서가 있거나, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서가 있다.",
            "3": "3. 무균제제가 아니다.",
            "4": "4. 포장 재질의 변경이 없다.",
            "5": "5. 액상제제(경구용 액제, 주사제, 점안제, 점이제 등), 폐에 적용하는 흡입제 및 반고형제제에 해당한다(유제 및 현탁제 제외)."
        }
    },
    "p3_13": {
        "title": "3.2.P.3 제조\n13. 비무균제제의 제조 규모 변경",
        "subitems": {
            "13a": "13a. 뱃치 생산규모(의동기준 제3조의2제5항의 대조약과 의약품동등성을 입증하거나 임상시험을 실시한 제제) 대비 10배수 이하의 변경",
            "13b": "13b. 일반 제제에서 10배 초과 생산 규모 변경",
            "13c": "13c. 제형의 특수성이 인정되는 제제 및 반고형제제에서 10배 초과의 생산규모 변경\n*제형의 특수성이 인정되는 제제 : 제제 기술의 변화로 인해 약품의 방출 또는 용출기전이 상이하여 체내 흡수량 또는 흡수속도의 변화를 확인할 필요가 있는 경우 등 제형의 특수성이 인정되는 제제(예;경피흡수제, 이식제, 서방형제제, 설하정, 정량용분무제제 중 폐에 적용하는 흡입제, 현탁성 주사제 등) 「의약품의 품목허가·신고·심사 규정 제25조2항제5호」"
        },
        "requirements": {
            "1": "1. 해당 변경은 제제의 재현성 및/또는 일관성에 영향을 미치지 않는다.",
            "2": "2. 제조장비의 작동원리 및 디자인이 동일하고, 제조 방법 및/또는 공정 중 관리의 변경은 제조 규모의 변경에 따른 변경만 필요로 한다.",
            "3": "3. (위험도 평가에 따른) 제조 공정 밸리데이션 실시 계획서가 있거나 또는 현재의 밸리데이션 실시 계획서에 따라 생산 규모의 3 배치에 대한 제조 공정 밸리데이션이 성공적으로 수행되었다.",
            "4": "4. 해당 변경은 제조 과정에서 발생하는 예기치 않은 사례 혹은 안정성 문제로 인하여 필요로 하는 것이 아니다.",
            "5": "5. 일반제제에 해당한다(장용성 및 방출조절제제, 서방성제제 등 제형의 특수성이 인정되는 제제 제외",
            "6": "6. 액상제제(경구용 액제, 외용액제, 점안제, 점이제 등)"
        }
    },
    "p3_14": {
        "title": "3.2.P.3 제조\n14. 무균제제의 제조 규모 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 해당 변경은 생산 일관성에 영향을 주지 않는다.",
            "2": "2. 원료약품 및 그 분량에는 변경이 없다.",
            "3": "3. 의약품의 규격 변경이 없다.",
            "4": "4. (위험도 평가에 따른) 밸리데이션 실시 계획서가 있거나 또는 현재의 밸리데이션 실시 계획서에 따라 생산 규모의 3배치에 대한 제조 밸리데이션이 성공적으로 수행되었다.",
            "5": "5. 액상제제(주사제, 점안제, 점이제 등)로서 제조 규모 변경이 10배 초과이다.",
            "6": "6. 제형의 특수성이 인정되는 제제,반고형제제 또는 분말형 주사제로서 제조 규모 변경이 10배 초과이다."
        }
    },
    "p3_15": {
        "title": "3.2.P.3 제조\n15. 완제의약품의 제조공정 변경",
        "subitems": {
            "15a": "15a. 완제의약품의 제조공정 변경",
            "15b": "15b. 완제의약품의 제조공정 변경(예 : 용출에 영향을 주는 첨가제 등급 변경, 결합액의 용매 양 변경, 용출에 영향을 주는 첨가제의 투입순서 변경, 코팅액의 용매 변경 등)",
            "15c": "15c. 완제의약품의 제조공정 변경(예 :결합액의 용매 종류 변경 등)"
        },
        "requirements": {
            "1": "1. 불순물 프로파일의 변경이 없고 물리 화학적 성질에 변경이 없다. 용출 프로파일은 대조약 배치의 프로파일과 유사하다.",
            "2": "2. 변경 전·후 제제의 제조공정은 동일한 원리이고 반제품은 동일하며, 공정에 사용되는 제조 용매에는 변경이 없다(습식 조립법에서 건식 조립법으로의 변경, 직접 분말 압축법에서 습식 또는 건식 과립 압축법으로의 변경, 또는 그 반대로의 변경은 제조 원리의 변경으로 간주된다).",
            "3": "3. 변경 전·후 작동원리가 같은 동일 계열의 제조장비를 사용한다.(제조장비의 작동원리 및 디자인의 동일 여부는 [부록 1], [부록 2]를 참고한다.)",
            "4": "4. 완제의약품의 품질에 영향을 미치는 제조공정(주요공정) 조건의 변동이 없다.",
            "5": "5. 반제품 또는 완제의약품의 규격 변경은 없다.",
            "6": "6. 제조 중에 발생하는 예기치 않은 사례로 인한 규격 미충족 또는 안정성 문제 때문에 발생한 변경은 아니다.",
            "7": "7. 해당 변경은 계량 및/또는 전달 기능을 제공하는 일차 포장과 관련된 포장이나 라벨링 공정을 포함하지 않는다.",
            "8": "8. 일반제제에 해당한다(장용성 및 방출조절제제, 서방성제제 등 제형의 특수성이 인정되는 제제 제외).",
            "9": "9. 액상제제(경구용 액제, 주사제, 점안제, 점이제 등)에 해당한다."
        }
    },
    "p3_16": {
        "title": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경",
        "subitems": {
            "16a": "16a. 공정관리시험 기준의 변경",
            "16b": "16b. 공정관리시험 기준의 변경",
            "16c": "16c. 공정관리시험 항목의 삭제",
            "16d": "16d. 새로운 공정관리시험 기준의 추가",
            "16e": "16e. 공정관리시험 방법의 변경"
        },
        "requirements": {
            "1": "1. 해당 변경은 공정관리 기준(예, 마손도, 경도, 입도, 밀도, 수분 등) 범위 내에 있다.",
            "2": "2. 해당 변경은 제조 중에 발생하는 예기치 않은 사례로 인한 규격 미충족 또는 안정성 문제로 인하여 필요로 하는 것이 아니다.",
            "3": "3. 삭제된 공정관리 항목은 불필요하거나 생략 가능한 것으로 입증되었으며, 제제의 주요 품질 특성(예:혼합균일성, 질량편차)에 미치는 영향이 없거나 적다.",
            "4": "4. 시험방법에는 변경이 없다."
        }
    },
    "p4_17": {
        "title": "3.2.P.4 첨가제의 관리\n17. 첨가제 기원의 변경",
        "subitems": {
            "17a": "17a. 동물성 기원에서 식물 또는 합성 기원의 원료로 변경",
            "17b": "17b. 식물 또는 합성 기원에서 동물성 기원, 또는 동물성 원료의 다른 동물성 기원으로 변경"
        },
        "requirements": {
            "1": "1. 첨가제와 완제의약품의 규격에는 변경이 없다."
        }
    },
    "p4_18": {
        "title": "3.2.P.4 첨가제의 관리\n18. 별규에 해당하는 첨가제의 규격 또는 시험방법 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 해당 변경은 제조 과정에서 발생하는 예기치 않은 사례로 인한 규격 미충족 혹은 안정성 문제 때문에 발생한 변경은 아니다."
        }
    },
    "p4_19": {
        "title": "3.2.P.4 첨가제의 관리\n19. 식약처장이 인정하는 공정서 규격으로 첨가제 규격의 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 공정서를 준수하기 위해 요구되는 규격 이외에는 변경이 없다(예: 입자 크기 분포의 변경 없음)."
        }
    },
    "p7_20": {
        "title": "3.2.P.7 용기-마개 시스템\n20. 비무균제제의 직접용기 및 포장 재질, 종류 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 고형제제로서 주성분, 제형, 투여경로가 동일한 기허가 의약품에서 사용례가 확인되는 용기 재질로의 변경 또는 용기포장 종류는 동일하고 보호성이 동등이상인 재질로의 변경으로, 기허가의약품의 사용(유효)기간을 초과하지 않는다."
        }
    },
    "p7_21": {
        "title": "3.2.P.7 용기-마개 시스템\n21. 무균제제의 직접용기 및 포장 재질, 종류 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 변경 후 용기의 보호성 등이 동등 이상이고 상호작용 위험이 없다."
        }
    },
    "p7_22": {
        "title": "3.2.P.7 용기-마개 시스템\n22. 직접 포장의 규격 변경",
        "subitems": {
            "22a": "22a. 규격 기준의 강화",
            "22b": "22b. 시험 항목의 추가 또는 삭제"
        },
        "requirements": {
            "1": "1. 모든 변경 사항은 현재의 허용 기준범위 내에 있다.",
            "2": "2. 해당 변경은 제조 과정에서 발생하는 예기치 않은 사례로 인한 규격 미충족 혹은 안정성 문제 때문에 발생한 변경은 아니다."
        }
    },
    "p7_23": {
        "title": "3.2.P.7 용기-마개 시스템\n23. 포장단위 변경",
        "subitems": {},
        "requirements": {
            "1": "1. 해당 변경은 제조 과정에서 발생하는 예기치 않은 사례로 인한 규격 미충족 혹은 안정성 문제 때문에 발생한 변경은 아니다.",
            "2": "2. 이외 허가사항의 변경은 없다."
        }
    },
    "ds_24": {
        "title": "디자인스페이스(Design Space)\n24. 새로운 디자인스페이스 도입 또는 허가된 디자인스페이스의 확장",
        "subitems": {
            "24a": "24a. 완제의약품/원료의약품 제조공정에서 하나 이상의 단위 조작 변경(공정 중 관리 및/또는 시험방법 포함)",
            "24b": "24b. 첨가제/반제품 및/또는 완제의약품/원료의약품의 시험방법 변경"
        },
    }
}

st.session_state["step6_items"] = step6_items

if "step6_selections" not in st.session_state:
    st.session_state.step6_selections = {}

if "step6_page" not in st.session_state:
    st.session_state.step6_page = 0

# ===== Step7 상태 초기화 =====
if "step7_page" not in st.session_state:
    st.session_state.step7_page = 0
if "step7_results" not in st.session_state:
    st.session_state.step7_results = {}

def go_to_prev_step6_page():
    if st.session_state.step6_page > 0:
        st.session_state.step6_page -= 1

def go_to_next_step6_page():
    if st.session_state.step6_page < len(st.session_state.step6_targets) - 1:
        st.session_state.step6_page += 1

def go_prev_step7_page():
    if st.session_state.step7_page > 0:
        st.session_state.step7_page -= 1
    else:
        st.session_state.step = 6

def go_next_step7_page():
    if st.session_state.step7_page < len(st.session_state.step6_targets) - 1:
        st.session_state.step7_page += 1

def go_to_step8():
    st.session_state.step8_page = 0
    st.session_state.step = 8

def go_back_to_step5():
    st.session_state.step = 5

def go_to_step7():
    st.session_state.step = 7

if st.session_state.step == 6:
    st.markdown(
        "<h2 style='font-size:25px; text-align:center; background-color:#FAF3DB; padding:4px;'>허가 후 제조방법 변경사항의 충족조건 선택</h2>",
        unsafe_allow_html=True,
    )
    # Add a blank line below the page title for better spacing
    st.write("")
    st.markdown(
        "<p style='font-size:16px;'>➤ 변경사항의 충족조건에 대한 충족 여부를 선택하세요.</p>",
        unsafe_allow_html=True,
    )


    st.markdown(
        """
        <style>
        div[data-testid="stRadio"] > label p {
            font-size:16px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stRadio"] label div[data-testid="stMarkdownContainer"] p {
            font-size:16px;
            margin-top:0;
            margin-bottom:0;
        }
        div[data-testid="stMarkdownContainer"] p {
            margin-top:0;
            margin-bottom:0;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    targets = st.session_state.step6_targets
    total_pages = len(targets)
    if not targets:
        st.warning("이전 단계에서 선택된 항목이 없습니다.")
    else:
        current_key = targets[st.session_state.step6_page]
        block = step6_items.get(current_key)

        if block:
            title_parts = block["title"].split("\n", 1)
            st.markdown(
                f"<p style='font-size:22px;'>{title_parts[0]}</p>",
                unsafe_allow_html=True,
            )
            if len(title_parts) > 1:
                st.markdown(
                    f"<p style='font-size:18px;'>{title_parts[1]}</p>",
                    unsafe_allow_html=True,
                )

            # Add a blank line before listing subitems or requirements
            st.write("")

            # 자동 동기화 쌍 정의
            sync_pairs = {
                "12c1": "12c2",
                "12c2": "12c1",
                "16a": "16b",
                "16b": "16a",
            }

            # 하위항목
            for sub_key, sub_text in block.get("subitems", {}).items():
                full_key = f"{current_key}_sub_{sub_key}"

                sub_label = sub_text
                st.markdown(
                    f"<p style='font-size:16px; margin-bottom:10px;'>{sub_label}</p>",
                    unsafe_allow_html=True,
                )

                if current_key == "p3_15":
                    st.session_state.step6_selections[full_key] = "변경 있음"
                    st.radio(
                        "",
                        ["변경 있음"],
                        index=0,
                        key=full_key,
                        disabled=True,
                        label_visibility="collapsed",
                    )

                elif sub_key in sync_pairs:
                    other = sync_pairs[sub_key]
                    other_key = f"{current_key}_sub_{other}"

                    current_value = st.session_state.step6_selections.get(full_key, "변경 없음")
                    current_value = st.radio(
                        "",
                        ["변경 있음", "변경 없음"],
                        key=full_key,
                        index=0 if current_value == "변경 있음" else 1,
                        label_visibility="collapsed",
                    )

                    st.session_state.step6_selections[full_key] = current_value
                    st.session_state.step6_selections[other_key] = current_value

                else:
                    st.session_state.step6_selections[full_key] = st.radio(
                        "",
                        ["변경 있음", "변경 없음"],
                        key=full_key,
                        label_visibility="collapsed",
                    )

            # 충족요건
            for req_key, req_text in block.get("requirements", {}).items():
                full_key = f"{current_key}_req_{req_key}"
                req_label = req_text.replace("\n", "<br>")
                st.markdown(
                    f"<p style='font-size:16px; margin-bottom:10px;'>{req_label}</p>",
                    unsafe_allow_html=True,
                )
                st.session_state.step6_selections[full_key] = st.radio(
                    "",
                    ["충족", "미충족"],
                    key=full_key,
                    label_visibility="collapsed",
                )

        else:
            st.warning("해당 항목 정보를 찾을 수 없습니다.")

    # 페이지 번호 표시 및 하단 네비게이션 영역은 targets 유무와 상관없이 노출
    st.markdown(
        f"<h6 style='text-align:center'>{(st.session_state.step6_page + 1) if targets else 0} / {total_pages}</h6>",
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.button(
            "이전단계로",
            on_click=go_back_to_step5 if st.session_state.step6_page == 0 else go_to_prev_step6_page,
        )
    with col2:
        if targets and st.session_state.step6_page == len(targets) - 1:
            st.button("결과 확인하기", on_click=go_to_step7)
        elif targets:
            st.button("다음항목 선택하기", on_click=go_to_next_step6_page)
        else:
            st.button("다음항목 선택하기", disabled=True)

# ===== Step7 상수 정의 =====
STEP7_ROWS = [
    {
        "title_key": "s1_1",
        "title_text": "3.2.S.1 일반정보\n1. 원료의약품 명칭변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s1_1_req_1\") == \"충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (S.1.1) 공정서 또는 국제 의약품 일반명 리스트(INN, The International Nonproprietary Name) 등 근거서류.\n2. 개정된 제품정보"
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2a\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_3\") == \"충족\" and step6_selections.get(\"s2_2_req_4\") == \"충족\" and step6_selections.get(\"s2_2_req_9\") == \"충족\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
         "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n10. 변경 전·후 출발 물질 또는 중간체의 최소 1배치에 대한 시험 성적서(해당하는 경우), 출발물질 또는 중간체 변경 전·후 최종 원료의약품 2배치에 대한 배치분석 자료."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2a\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_3\") == \"미충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n3. (S.2.5) 무균원료의약품 생산의 경우 (위험도 평가 결과에 따른) 무균공정에 대한 공정밸리데이션 자료 및 평가 자료.\n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n7. (S.4.1) 원료의약품 기준 및 시험방법에 관한 자료.\n10. 변경 전·후 출발 물질 또는 중간체의 최소 1배치에 대한 시험 성적서(해당하는 경우), 출발물질 또는 중간체 변경 전·후 최종 원료의약품 2배치에 대한 배치분석 자료.\n11. (S.7.2) 변경 후 원료의약품의 안정성 시험 필요성 고찰 및 필요한 경우 안정성 시험 이행 계획서."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2b\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_3\") == \"충족\" and step6_selections.get(\"s2_2_req_5\") == \"충족\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n8. (S.2) 원료의약품 및 핵심(최종) 중간체(해당되는 경우) 합성 경로, 사용 원료, 품질 관리 절차 및 규격 변경이 없다는 확인서(statement).\n10. 변경 전·후 출발 물질 또는 중간체의 최소 1배치에 대한 시험 성적서(해당하는 경우), 출발물질 또는 중간체 변경 전·후 최종 원료의약품 2배치에 대한 배치분석 자료."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2b\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_3\") == \"미충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n3. (S.2.5) 무균원료의약품 생산의 경우 (위험도 평가 결과에 따른) 무균공정에 대한 공정밸리데이션 자료 및 평가 자료.\n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n7. (S.4.1) 원료의약품 기준 및 시험방법에 관한 자료.\n10. 변경 전·후 출발 물질 또는 중간체의 최소 1배치에 대한 시험 성적서(해당하는 경우), 출발물질 또는 중간체 변경 전·후 최종 원료의약품 2배치에 대한 배치분석 자료.\n11. (S.7.2) 변경 후 원료의약품의 안정성 시험 필요성 고찰 및 필요한 경우 안정성 시험 이행 계획서."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2c\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_1\") == \"충족\" and step6_selections.get(\"s2_2_req_2\") == \"충족\" and step6_selections.get(\"s2_2_req_3\") == \"충족\" and step6_selections.get(\"s2_2_req_6\") == \"충족\" and step6_selections.get(\"s2_2_req_7\") == \"충족\" and step6_selections.get(\"s2_2_req_8\") == \"충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n5. (S.4.4) 변경 전·후 원료의약품 2배치(파일럿 배치 이상)에 대한 배치분석자료 .\n8. (S.2) 원료의약품 및 핵심(최종) 중간체(해당되는 경우) 합성 경로, 사용 원료, 품질 관리 절차 및 규격 변경이 없다는 확인서(statement)."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2c\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_3\") == \"미충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\" and step6_selections.get(\"s2_2_req_10\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당되는 경우) 해당 품목을 제조하는 제조소에 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서, 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서. \n2. (S.2.1) 제조소명, 주소, 책임부과범위 및 해당하는 경우 수탁업소에 관한 자료. \n3. (S.2.5) 무균원료의약품 생산의 경우 (위험도 평가 결과에 따른) 무균공정에 대한 공정밸리데이션 자료 및 평가 자료.\n4. 변경 전·후 제조소의 원료의약품, 중간체 또는 원료의약품 출발 물질 (해당되는 경우)제조 공정에 관한 자료.\n5. (S.4.4) 변경 전·후 원료의약품 2배치(파일럿 배치 이상)에 대한 배치분석자료 .\n6. (P.8.2) 원료의약품의 품질 특성이 완제의약품의 안정성에 영향을 미칠 수 있는 변경의 경우, 완제의약품 1배치(실제 생산 규모의)에 대한 안정성 시험 이행 계획서.\n7. (S.4.1) 원료의약품 기준 및 시험방법에 관한 자료.\n9. 변경 후 원료의약품이 완제의약품의 안전성, 유효성 및 품질에 미치는 영향에 대한 고찰자료.\n10. 변경 전·후 출발 물질 또는 중간체의 최소 1배치에 대한 시험 성적서(해당하는 경우), 출발물질 또는 중간체 변경 전·후 최종 원료의약품 2배치에 대한 배치분석 자료.\n11. (S.7.2) 변경 후 원료의약품의 안정성 시험 필요성 고찰 및 필요한 경우 안정성 시험 이행 계획서."
    },
    {
        "title_key": "s2_2",
        "title_text": "3.2.S.2 제조\n2. 원료의약품의 제조소 또는 제조업자의 변경 또는 추가\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_2_sub_2c\") == \"변경 있음\" and step6_selections.get(\"s2_2_req_10\") == \"충족\" and step6_selections.get(\"s2_2_req_1\") == \"미충족\" and step6_selections.get(\"s2_2_req_2\") == \"미충족\" and step6_selections.get(\"s2_2_req_3\") == \"미충족\" and step6_selections.get(\"s2_2_req_4\") == \"미충족\" and step6_selections.get(\"s2_2_req_5\") == \"미충족\" and step6_selections.get(\"s2_2_req_6\") == \"미충족\" and step6_selections.get(\"s2_2_req_7\") == \"미충족\" and step6_selections.get(\"s2_2_req_8\") == \"미충족\" and step6_selections.get(\"s2_2_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n12. 해당국가의 공식기관에서 발급받은 문서(GMP 증명서 등 포함) 또는 인증여부를 확인할 수 있는 자료(제조소의 책임자가 서명하고 공증받은 자료 등). "
    },
    {
        "title_key": "s2_3",
        "title_text": "3.2.S.2 제조\n3. 원료의약품 제조 공정의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_3_req_1\") == \"충족\" and step6_selections.get(\"s2_3_req_2\") == \"충족\" and step6_selections.get(\"s2_3_req_3\") == \"충족\" and step6_selections.get(\"s2_3_req_4\") == \"충족\" and step6_selections.get(\"s2_3_req_5\") == \"충족\" and step6_selections.get(\"s2_3_req_6\") == \"충족\" and step6_selections.get(\"s2_3_req_7\") == \"충족\" and step6_selections.get(\"s2_3_req_8\") == \"충족\" and step6_selections.get(\"s2_3_req_9\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. 변경 전·후 제조방법 비교표 등 변경 전·후에 관한 자료\n3. (S.2.2) 변경하고자 하는 합성 공정 흐름도 및 상세 제조 공정에 관한 자료.\n4. (S.2.3)(해당되는 경우) 변경하고자 하는 원료의약품 제조에 사용된 원료(예 : 원료약품, 출발 물질, 용매, 시약, 촉매)의 규격 및 시험 성적서.\n11. (S.4.4) 변경 전·후 원료의약품 최소 2배치(파일럿 배치 이상)에 대한 배치분석 자료."
    },
    {
        "title_key": "s2_3",
        "title_text": "3.2.S.2 제조\n3. 원료의약품 제조 공정의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_3_req_1\") == \"충족\" and step6_selections.get(\"s2_3_req_2\") == \"충족\" and step6_selections.get(\"s2_3_req_3\") == \"충족\" and step6_selections.get(\"s2_3_req_5\") == \"충족\" and step6_selections.get(\"s2_3_req_6\") == \"충족\" and step6_selections.get(\"s2_3_req_7\") == \"충족\" and step6_selections.get(\"s2_3_req_9\") == \"충족\" and step6_selections.get(\"s2_3_req_4\") == \"미충족\" and step6_selections.get(\"s2_3_req_8\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. 변경 전·후 제조방법 비교표 등 변경 전·후에 관한 자료\n3. (S.2.2) 변경하고자 하는 합성 공정 흐름도 및 상세 제조 공정에 관한 자료.\n10. (S.4.1) 변경 후 원료의약품 기준 및 시험방법에 관한 자료.(변경되는 경우, 출발물질 및 중간체의 기준 및 시험방법)\n11. (S.4.4) 변경 전·후 원료의약품 최소 2배치(파일럿 배치 이상)에 대한 배치분석 자료."
    },
    {
        "title_key": "s2_3",
        "title_text": "3.2.S.2 제조\n3. 원료의약품 제조 공정의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_3_req_1\") == \"충족\" and step6_selections.get(\"s2_3_req_2\") == \"충족\" and step6_selections.get(\"s2_3_req_3\") == \"충족\" and step6_selections.get(\"s2_3_req_4\") == \"충족\" and step6_selections.get(\"s2_3_req_5\") == \"충족\" and step6_selections.get(\"s2_3_req_6\") == \"충족\" and step6_selections.get(\"s2_3_req_7\") == \"미충족\" and step6_selections.get(\"s2_3_req_8\") == \"미충족\" and step6_selections.get(\"s2_3_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. 변경 전·후 제조방법 비교표 등 변경 전·후에 관한 자료\n3. (S.2.2) 변경하고자 하는 합성 공정 흐름도 및 상세 제조 공정에 관한 자료.\n10. (S.4.1) 변경 후 원료의약품 기준 및 시험방법에 관한 자료.(변경되는 경우, 출발물질 및 중간체의 기준 및 시험방법)\n11. (S.4.4) 변경 전·후 원료의약품 최소 2배치(파일럿 배치 이상)에 대한 배치분석 자료."
    },
    {
        "title_key": "s2_3",
        "title_text": "3.2.S.2 제조\n3. 원료의약품 제조 공정의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_3_req_1\") == \"미충족\" and step6_selections.get(\"s2_3_req_2\") == \"미충족\" and step6_selections.get(\"s2_3_req_3\") == \"미충족\" and step6_selections.get(\"s2_3_req_4\") == \"미충족\" and step6_selections.get(\"s2_3_req_5\") == \"미충족\" and step6_selections.get(\"s2_3_req_6\") == \"미충족\" and step6_selections.get(\"s2_3_req_7\") == \"미충족\" and step6_selections.get(\"s2_3_req_8\") == \"미충족\" and step6_selections.get(\"s2_3_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 후 원료의약품이 완제의약품의 안전성, 유효성 및 품질에 미치는 영향에 대한 고찰자료. (완제의약품의 불순물 프로파일, 배치분석자료, 필요한 경우 안정성 자료 등)\n2. 변경 전·후 제조방법 비교표 등 변경 전·후에 관한 자료\n3. (S.2.2) 변경하고자 하는 합성 공정 흐름도 및 상세 제조 공정에 관한 자료.\n4. (S.2.3)(해당되는 경우) 변경하고자 하는 원료의약품 제조에 사용된 원료(예 : 원료약품, 출발 물질, 용매, 시약, 촉매)의 규격 및 시험 성적서.\n5. (S.2.3) 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위험 적합성 평가 자료.\n6. (S.2.4)(해당되는 경우) 주요 공정 및 중간체 관리에 관한 자료.\n7. (S.2.5) (위험도 평가에 따른) 멸균 공정 밸리데이션 자료 또는 멸균 평가 시험 자료.\n8. (S.3.1) 원료의약품의 구조 결정 자료(IR, UV 등) 및 물리화학적 성질에 관한 자료.\n9. (S.3.2) 불순물에 대한 고찰 및 근거자료.\n10. (S.4.1) 변경 후 원료의약품 기준 및 시험방법에 관한 자료.(변경되는 경우, 출발물질 및 중간체의 기준 및 시험방법)\n11. (S.4.4) 변경 전·후 원료의약품 최소 2배치(파일럿 배치 이상)에 대한 배치분석 자료.\n12. (S.7.1) 변경 후 원료의약품 최소 2배치(파일럿 배치 이상)에 대한 3개월 이상 가속 시험 (필요한 경우, 중간조건시험) 및 장기 보존 시험 자료.\n13. 난용성 원료의약품의 결정형 또는 입자도의 변경이 발생한 경우, 완제의약품의 품질과 생체이용률에 영향을 미치지 않는다는 입증 자료."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4a\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"충족\" and step6_selections.get(\"s2_4_req_2\") == \"충족\" and step6_selections.get(\"s2_4_req_3\") == \"충족\" and step6_selections.get(\"s2_4_req_4\") == \"미충족\" and step6_selections.get(\"s2_4_req_5\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4b\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"충족\" and step6_selections.get(\"s2_4_req_2\") == \"미충족\" and step6_selections.get(\"s2_4_req_3\") == \"미충족\" and step6_selections.get(\"s2_4_req_4\") == \"미충족\" and step6_selections.get(\"s2_4_req_5\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료.\n4. 해당되는 경우, 분석 방법 상세 자료.\n5. 공정관리 시험(추가, 교체, 삭제, 완화되는) 규격에 대한 타당성 입증 자료 또는 설명자료."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4c\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"미충족\" and step6_selections.get(\"s2_4_req_2\") == \"미충족\" and step6_selections.get(\"s2_4_req_3\") == \"미충족\" and step6_selections.get(\"s2_4_req_4\") == \"미충족\" and step6_selections.get(\"s2_4_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
     "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료.\n5. 공정관리 시험(추가, 교체, 삭제, 완화되는) 규격에 대한 타당성 입증 자료 또는 설명자료.\n7. (S.2.5) 무균원료의약품인 경우, 해당 공정기준 변경이 제품의 무균 및 멸균공정에 영향을 미칠 경우(위험도 평가 결과에 따른), 해당 공정에 대한 밸리데이션 자료 또는 평가 자료.\n8. (S.3.2) 해당변경이 불순물에 영향을 미칠 경우, 불순물에 대한 고찰 및 근거자료.\n9. (S.4.1) 변경 후 원료의약품(해당되는 경우 중간체) 규격에 관한 자료.\n10. (S.4.4) 변경 전·후 원료의약품 최소 1배치(파일럿 배치 이상)의 배치분석자료."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4d\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"충족\" and step6_selections.get(\"s2_4_req_4\") == \"충족\" and step6_selections.get(\"s2_4_req_5\") == \"충족\" and step6_selections.get(\"s2_4_req_2\") == \"미충족\" and step6_selections.get(\"s2_4_req_3\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료.\n6. 삭제되는 공정관리시험이 품질에 영향을 미치지 않음을 입증하는 자료(또는 위험 평가 자료)."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4d\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"충족\" and step6_selections.get(\"s2_4_req_2\") == \"미충족\" and step6_selections.get(\"s2_4_req_3\") == \"미충족\" and step6_selections.get(\"s2_4_req_4\") == \"미충족\" and step6_selections.get(\"s2_4_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료.\n5. 공정관리 시험(추가, 교체, 삭제, 완화되는) 규격에 대한 타당성 입증 자료 또는 설명자료.\n7. (S.2.5) 무균원료의약품인 경우, 해당 공정기준 변경이 제품의 무균 및 멸균공정에 영향을 미칠 경우(위험도 평가 결과에 따른), 해당 공정에 대한 밸리데이션 자료 또는 평가 자료.\n8. (S.3.2) 해당변경이 불순물에 영향을 미칠 경우, 불순물에 대한 고찰 및 근거자료.\n9. (S.4.1) 변경 후 원료의약품(해당되는 경우 중간체) 규격에 관한 자료.\n10. (S.4.4) 변경 전·후 원료의약품 최소 1배치(파일럿 배치 이상)의 배치분석자료."
    },
    {
        "title_key": "s2_4",
        "title_text": "3.2.S.2 제조\n4. 원료의약품 제조 공정관리 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_4_sub_4e\") == \"변경 있음\" and step6_selections.get(\"s2_4_req_1\") == \"충족\" and step6_selections.get(\"s2_4_req_2\") == \"미충족\" and step6_selections.get(\"s2_4_req_3\") == \"미충족\" and step6_selections.get(\"s2_4_req_4\") == \"미충족\" and step6_selections.get(\"s2_4_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
         "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 공정관리 시험 비교표 등 변경 전·후에 관한 자료\n2. (S.2.2) 변경 후 합성 공정 흐름도 및 제조 공정에 대한 서술 자료.\n3. (S.2.4) 변경 후 공정관리 시험 규격에 관한 자료.\n5. 공정관리 시험(추가, 교체, 삭제, 완화되는) 규격에 대한 타당성 입증 자료 또는 설명자료.\n7. (S.2.5) 무균원료의약품인 경우, 해당 공정기준 변경이 제품의 무균 및 멸균공정에 영향을 미칠 경우(위험도 평가 결과에 따른), 해당 공정에 대한 밸리데이션 자료 또는 평가 자료.\n8. (S.3.2) 해당변경이 불순물에 영향을 미칠 경우, 불순물에 대한 고찰 및 근거자료.\n9. (S.4.1) 변경 후 원료의약품(해당되는 경우 중간체) 규격에 관한 자료.\n10. (S.4.4) 변경 전·후 원료의약품 최소 1배치(파일럿 배치 이상)의 배치분석자료."
    },
    {
        "title_key": "s2_5",
        "title_text": "3.2.S.2 제조\n5. 원료의약품 또는 중간체의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_5_sub_5a\") == \"변경 있음\" and step6_selections.get(\"s2_5_req_1\") == \"충족\" and step6_selections.get(\"s2_5_req_2\") == \"충족\" and step6_selections.get(\"s2_5_req_3\") == \"충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (S.2.2) 변경 전·후 제조공정 비교표 및 변경 후 상세 제조 공정에 관한 자료.\n2. (S.2.5) (해당되는 경우, 위험도 평가에 따른) 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (S.4.1) 원료의약품의 규격에 관한 자료(해당되는 경우 중간체 규격에 관한 자료).\n4. (S.4.4) 변경 전·후 제조 규모에서 각 최소 1배치에 대한 배치 분석 자료."
    },
    {
        "title_key": "s2_5",
        "title_text": "3.2.S.2 제조\n5. 원료의약품 또는 중간체의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_5_sub_5b\") == \"변경 있음\" and step6_selections.get(\"s2_5_req_1\") == \"충족\" and step6_selections.get(\"s2_5_req_2\") == \"미충족\" and step6_selections.get(\"s2_5_req_3\") == \"충족\"\n)",
        "output_1_tag": "AR",
         "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (S.2.2) 변경 전·후 제조공정 비교표 및 변경 후 상세 제조 공정에 관한 자료.\n2. (S.2.5) (해당되는 경우, 위험도 평가에 따른) 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (S.4.1) 원료의약품의 규격에 관한 자료(해당되는 경우 중간체 규격에 관한 자료).\n4. (S.4.4) 변경 전·후 제조 규모에서 각 최소 1배치에 대한 배치 분석 자료."
    },
    {
        "title_key": "s2_5",
        "title_text": "3.2.S.2 제조\n5. 원료의약품 또는 중간체의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_5_sub_5c\") == \"변경 있음\" and step6_selections.get(\"s2_5_req_1\") == \"충족\" and step6_selections.get(\"s2_5_req_2\") == \"충족\" and step6_selections.get(\"s2_5_req_3\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (S.2.2) 변경 전·후 제조공정 비교표 및 변경 후 상세 제조 공정에 관한 자료.\n2. (S.2.5) (해당되는 경우, 위험도 평가에 따른) 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (S.4.1) 원료의약품의 규격에 관한 자료(해당되는 경우 중간체 규격에 관한 자료).\n5. (S.4.4) 변경 전·후 제조 규모에서 각 최소 2배치에 대한 배치 분석 자료."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6a\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_1\") == \"충족\" and step6_selections.get(\"s2_6_req_2\") == \"충족\" and step6_selections.get(\"s2_6_req_3\") == \"충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_5\") == \"미충족\" and step6_selections.get(\"s2_6_req_6\") == \"미충족\" and step6_selections.get(\"s2_6_req_7\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\" and step6_selections.get(\"s2_6_req_9\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서)."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6b\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_4\") == \"충족\" and step6_selections.get(\"s2_6_req_5\") == \"충족\" and step6_selections.get(\"s2_6_req_6\") == \"충족\" and step6_selections.get(\"s2_6_req_1\") == \"미충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_3\") == \"미충족\" and step6_selections.get(\"s2_6_req_7\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\" and step6_selections.get(\"s2_6_req_9\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서)."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6c\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_1\") == \"충족\" and step6_selections.get(\"s2_6_req_5\") == \"충족\" and step6_selections.get(\"s2_6_req_6\") == \"충족\" and step6_selections.get(\"s2_6_req_7\") == \"충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_3\") == \"미충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\" and step6_selections.get(\"s2_6_req_9\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서)."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6d\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_1\") == \"충족\" and step6_selections.get(\"s2_6_req_8\") == \"충족\" and step6_selections.get(\"s2_6_req_9\") == \"충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_3\") == \"미충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_5\") == \"미충족\" and step6_selections.get(\"s2_6_req_6\") == \"미충족\" and step6_selections.get(\"s2_6_req_7\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서).\n4. 해당 변경이 품질에 영향을 미치지 않음을 입증하는 자료(또는 위험 평가 자료)."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6e\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_1\") == \"미충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_3\") == \"미충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_5\") == \"미충족\" and step6_selections.get(\"s2_6_req_6\") == \"미충족\" and step6_selections.get(\"s2_6_req_7\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\" and step6_selections.get(\"s2_6_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서).\n4. 해당 변경이 품질에 영향을 미치지 않음을 입증하는 자료(또는 위험 평가 자료).\n5. (S.3.2) (해당하는 경우) 불순물에 대한 고찰 및 근거자료."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6f\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_3\") == \"충족\" and step6_selections.get(\"s2_6_req_6\") == \"충족\" and step6_selections.get(\"s2_6_req_7\") == \"충족\" and step6_selections.get(\"s2_6_req_9\") == \"충족\" and step6_selections.get(\"s2_6_req_1\") == \"미충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_5\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서)."
    },
    {
        "title_key": "s2_6",
        "title_text": "3.2.S.2 제조\n6. 원료의약품의 제조에 사용되는 원료(출발물질, 중간체, 용매, 시약 등)의 규격변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"s2_6_sub_6g\") == \"변경 있음\" and step6_selections.get(\"s2_6_req_1\") == \"미충족\" and step6_selections.get(\"s2_6_req_2\") == \"미충족\" and step6_selections.get(\"s2_6_req_3\") == \"미충족\" and step6_selections.get(\"s2_6_req_4\") == \"미충족\" and step6_selections.get(\"s2_6_req_5\") == \"미충족\" and step6_selections.get(\"s2_6_req_6\") == \"미충족\" and step6_selections.get(\"s2_6_req_7\") == \"미충족\" and step6_selections.get(\"s2_6_req_8\") == \"미충족\" and step6_selections.get(\"s2_6_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료.\n2. (S.2.3) 원료의약품의 제조에 사용하는 변경된 원료의 정보(규격 또는 공급처 성적서).\n3. (S.2.4) (해당하는 경우) 변경된 중간체에 대한 정보(규격 또는 공급처 성적서).\n4. 해당 변경이 품질에 영향을 미치지 않음을 입증하는 자료(또는 위험 평가 자료).\n5. (S.3.2) (해당하는 경우) 불순물에 대한 고찰 및 근거자료."
    },
    {
        "title_key": "p1_7",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n7. 완제의약품 중 고형 제제의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_7_sub_7a\") == \"변경 있음\" and step6_selections.get(\"p1_7_req_1\") == \"충족\" and step6_selections.get(\"p1_7_req_4\") == \"충족\" and step6_selections.get(\"p1_7_req_2\") == \"미충족\" and step6_selections.get(\"p1_7_req_3\") == \"미충족\" and step6_selections.get(\"p1_7_req_5\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.1) 완제의약품의 성상 및 원료약품 분량.\n3. (P.2) 변경하고자 하는 제제의 성분에 대한 검토 자료(예: 첨가제의 선택, 원료의약품과 첨가제의 배합 적합성).\n4. (P.3) 배치 조성에 대한 자료. \n5. (P.4) 첨가제 종류를 변경/추가하는 경우, 첨가제의 규격에 관한 자료.\n7. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상)에 대한 시험 성적서. \n10. (P.8.2) 변경하고자 하는 제제의 생산 규모 배치에 대한 안정성 시험 계획서 및 이행서약."
    },
    {
        "title_key": "p1_7",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n7. 완제의약품 중 고형 제제의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_7_sub_7b\") == \"변경 있음\" and step6_selections.get(\"p1_7_req_1\") == \"충족\" and step6_selections.get(\"p1_7_req_2\") == \"충족\" and step6_selections.get(\"p1_7_req_3\") == \"충족\" and step6_selections.get(\"p1_7_req_4\") == \"충족\" and step6_selections.get(\"p1_7_req_5\") == \"충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.1) 완제의약품의 성상 및 원료약품 분량.\n5. (P.4) 첨가제 종류를 변경/추가하는 경우, 첨가제의 규격에 관한 자료.\n7. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상)에 대한 시험 성적서. \n8. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료.\n10. (P.8.2) 변경하고자 하는 제제의 생산 규모 배치에 대한 안정성 시험 계획서 및 이행서약."
    },
    {
        "title_key": "p1_7",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n7. 완제의약품 중 고형 제제의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_7_sub_7b\") == \"변경 있음\" and step6_selections.get(\"p1_7_req_1\") == \"충족\" and step6_selections.get(\"p1_7_req_2\") == \"충족\" and step6_selections.get(\"p1_7_req_3\") == \"충족\" and step6_selections.get(\"p1_7_req_4\") == \"충족\" and step6_selections.get(\"p1_7_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.1) 완제의약품의 성상 및 원료약품 분량.\n5. (P.4) 첨가제 종류를 변경/추가하는 경우, 첨가제의 규격에 관한 자료.\n7. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상)에 대한 시험 성적서. \n8. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료."
    },
    {
        "title_key": "p1_7",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n7. 완제의약품 중 고형 제제의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_7_sub_7c\") == \"변경 있음\" and step6_selections.get(\"p1_7_req_1\") == \"충족\" and step6_selections.get(\"p1_7_req_4\") == \"충족\" and step6_selections.get(\"p1_7_req_2\") == \"미충족\" and step6_selections.get(\"p1_7_req_3\") == \"미충족\" and step6_selections.get(\"p1_7_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표2] 원료약품 및 그 분량 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.1) 완제의약품의 성상 및 원료약품 분량.\n3. (P.2) 변경하고자 하는 제제의 성분에 대한 검토 자료(예: 첨가제의 선택, 원료의약품과 첨가제의 배합 적합성).\n4. (P.3) 배치 조성에 대한 자료. \n5. (P.4) 첨가제 종류를 변경/추가하는 경우, 첨가제의 규격에 관한 자료.\n6. (P.4.5) 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위험 적합성 평가 자료.\n7. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상)에 대한 시험 성적서. \n8. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료.\n9. (P.8.1) 변경 후 완제의약품 최소 2배치(파일럿 배치 이상)에 대한 장기 및 가속 안정성 시험 최소 3개월 자료. 생물학적동등성시험을 제출하는 경우, 장기 및 가속 최소 6개월 자료.\n10. (P.8.2) 변경하고자 하는 제제의 생산 규모 배치에 대한 안정성 시험 계획서 및 이행서약."
    },
    {
        "title_key": "p1_8",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n8. 고형제제의 코팅층 무게 변경\n",
        "output_condition_all_met": "(\n    \n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.1) 완제의약품의 원료약품 분량.\n2. (P.2 또는 R) 「의약품동등성시험기준」 [별표2] 원료약품 및 그 분량 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n3. (P.3) 배치 조성에 대한 자료.\n4. (P.5) 최소 1배치(파일럿 배치 이상) 완제의약품 기준 및 시험방법 및 시험 성적서. "
    },
    {
        "title_key": "p1_9",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n9. 완제의약품(고형제제 제외)의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_9_req_9\") == \"충족\" and step6_selections.get(\"p1_9_req_8\") == \"미충족\" and (\n        step6_selections.get(\"p1_9_req_1\") == \"충족\" or step6_selections.get(\"p1_9_req_2\") == \"충족\" or step6_selections.get(\"p1_9_req_3\") == \"충족\" or step6_selections.get(\"p1_9_req_4\") == \"충족\" or step6_selections.get(\"p1_9_req_5\") == \"충족\" or step6_selections.get(\"p1_9_req_6\") == \"충족\" or step6_selections.get(\"p1_9_req_7\") == \"충족\"\n   )\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표2] 원료약품 및 그 분량 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료로서, 완제의약품의 물리화학적 특성에 변화가 없음을 입증할 수 있는 이화학적동등성시험자료(예. 점도, 삼투압, pH 등).\n3. (P.1) 완제의약품의 성상 및 조성.\n4. (P.2) 변경하고자 하는 제제의 구성 성분에 대한 검토 자료(예: 첨가제의 선택, 원료약품과 첨가제 간의 배합성, 변경된 제제의 포장 시스템 적합성 시험), 주사제, 점안제 및 점이제의 경우 유효성분에 영향을 미치지 않음을 입증하는 자료로서 변경가능 첨가제 내 종류가 상이하거나 종류는 동일하고 분량만 상이한 경우 가속 6개월 안정성 시험자료 제출\n5. (P.3) 배치 조성, 제조 공정 및 공정 관리의 설명 자료, 중요 공정 및 중간체의 관리, (위험도 평가에 따른) 공정 밸리데이션 계획서 및/또는 평가에 대한 자료.\n6. (P.4) 첨가제 종류를 변경/추가하는 경우, 해당 첨가제의 규격에 관한 자료.\n7. (P.4.5) 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위해 적합성 평가 자료.\n8. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상) 시험 성적서. \n9. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료.\n10. (P.8.1) 변경 후 완제의약품 최소 2배치(파일럿 배치 이상)에 대한 장기 및 가속 안정성 시험 최소 3개월 자료. 생물학적동등성시험을 제출하는 경우, 장기 및 가속 최소 6개월 자료\n11. (P.8.2) 변경하고자 하는 제제의 생산규모 배치에 대한 안정성 시험 계획서 및 이행서약."
    },
    {
        "title_key": "p1_9",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n9. 완제의약품(고형제제 제외)의 조성 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_9_req_8\") == \"충족\" and step6_selections.get(\"p1_9_req_9\") == \"충족\" and step6_selections.get(\"p1_9_req_1\") == \"미충족\" and step6_selections.get(\"p1_9_req_2\") == \"미충족\" and step6_selections.get(\"p1_9_req_3\") == \"미충족\" and step6_selections.get(\"p1_9_req_4\") == \"미충족\" and step6_selections.get(\"p1_9_req_5\") == \"미충족\" and step6_selections.get(\"p1_9_req_6\") == \"미충족\" and step6_selections.get(\"p1_9_req_7\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.2 또는 R) 「의약품동등성시험기준」 [별표2] 원료약품 및 그 분량 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료로서, 생물학적동등성시험자료 또는 비교임상시험자료, 또는 생물학적동등성시험자료를 갈음할 수 있는 타당한 자료.\n3. (P.1) 완제의약품의 성상 및 조성.\n4. (P.2) 변경하고자 하는 제제의 구성 성분에 대한 검토 자료(예: 첨가제의 선택, 원료약품과 첨가제 간의 배합성, 변경된 제제의 포장 시스템 적합성 시험), 주사제, 점안제 및 점이제의 경우 유효성분에 영향을 미치지 않음을 입증하는 자료로서 변경가능 첨가제 내 종류가 상이하거나 종류는 동일하고 분량만 상이한 경우 가속 6개월 안정성 시험자료 제출\n5. (P.3) 배치 조성, 제조 공정 및 공정 관리의 설명 자료, 중요 공정 및 중간체의 관리, (위험도 평가에 따른) 공정 밸리데이션 계획서 및/또는 평가에 대한 자료.\n6. (P.4) 첨가제 종류를 변경/추가하는 경우, 해당 첨가제의 규격에 관한 자료.\n7. (P.4.5) 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위해 적합성 평가 자료.\n8. (P.5) 완제의약품의 기준 및 시험방법, 최소 1배치(파일럿 배치 이상) 시험 성적서. \n9. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료.\n10. (P.8.1) 변경 후 완제의약품 최소 2배치(파일럿 배치 이상)에 대한 장기 및 가속 안정성 시험 최소 3개월 자료. 생물학적동등성시험을 제출하는 경우, 장기 및 가속 최소 6개월 자료\n11. (P.8.2) 변경하고자 하는 제제의 생산규모 배치에 대한 안정성 시험 계획서 및 이행서약."
    },
    {
        "title_key": "p1_10",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n10. 완제의약품(고형제제 제외)에 쓰이는 착색제 또는 착향제의 종류와 분량의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_10_sub_10a\") == \"변경 있음\" and step6_selections.get(\"p1_10_req_1\") == \"충족\" and step6_selections.get(\"p1_10_req_3\") == \"충족\" and step6_selections.get(\"p1_10_req_4\") == \"충족\" and step6_selections.get(\"p1_10_req_8\") == \"충족\" and step6_selections.get(\"p1_10_req_9\") == \"충족\" and step6_selections.get(\"p1_10_req_2\") == \"미충족\" and step6_selections.get(\"p1_10_req_5\") == \"미충족\" and step6_selections.get(\"p1_10_req_6\") == \"미충족\" and step6_selections.get(\"p1_10_req_7\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.1) 완제의약품의 원료약품 분량.\n2. (P.3) 배치 조성에 대한 자료.\n5. (P.5) 완제의약품 기준 및 시험방법 및 최소 1배치(파일럿 배치 이상) 이상의 시험 성적서.\n7. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n8. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p1_10",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n10. 완제의약품(고형제제 제외)에 쓰이는 착색제 또는 착향제의 종류와 분량의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_10_sub_10b\") == \"변경 있음\" and step6_selections.get(\"p1_10_req_1\") == \"충족\" and step6_selections.get(\"p1_10_req_2\") == \"충족\" and step6_selections.get(\"p1_10_req_3\") == \"충족\" and step6_selections.get(\"p1_10_req_4\") == \"충족\" and step6_selections.get(\"p1_10_req_8\") == \"충족\" and step6_selections.get(\"p1_10_req_9\") == \"충족\" and step6_selections.get(\"p1_10_req_5\") == \"미충족\" and step6_selections.get(\"p1_10_req_6\") == \"미충족\" and step6_selections.get(\"p1_10_req_7\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.1) 완제의약품의 원료약품 분량.\n2. (P.3) 배치 조성에 대한 자료.\n5. (P.5) 완제의약품 기준 및 시험방법 및 최소 1배치(파일럿 배치 이상) 이상의 시험 성적서.\n7. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n8. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p1_10",
        "title_text": "3.2.P.1 완제의약품의 성상 및 조성\n10. 완제의약품(고형제제 제외)에 쓰이는 착색제 또는 착향제의 종류와 분량의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p1_10_sub_10c\") == \"변경 있음\" and step6_selections.get(\"p1_10_req_1\") == \"충족\" and step6_selections.get(\"p1_10_req_2\") == \"충족\" and step6_selections.get(\"p1_10_req_3\") == \"충족\" and step6_selections.get(\"p1_10_req_5\") == \"충족\" and step6_selections.get(\"p1_10_req_6\") == \"충족\" and step6_selections.get(\"p1_10_req_7\") == \"충족\" and step6_selections.get(\"p1_10_req_8\") == \"충족\" and step6_selections.get(\"p1_10_req_9\") == \"충족\" and step6_selections.get(\"p1_10_req_4\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.1) 완제의약품의 원료약품 분량.\n2. (P.3) 배치 조성에 대한 자료.\n3. (P.2) 완제의약품의 성분에 대한 검토 자료(예: 착색제나 착향제가 관련 규격이 첨부된 혼합물로 구입된 경우, 원료의약품과 착색제나 착향제의 성분 조성과의 배합 적합성).\n4. (P.4) 착색제 또는 착향제의 규격 및 성적서. 사람 또는 동물 유래 물질이 사용되는 경우, 출처가 새로운 원료약품에 대한 BSE/TSE(소해면상뇌증/전염성해면상뇌증) 위해 적합성 평가 자료.\n5. (P.5) 완제의약품 기준 및 시험방법 및 최소 1배치(파일럿 배치 이상) 이상의 시험 성적서.\n6. (P.5.3) 해당되는 경우, 변경된 첨가제가 완제의약품의 분석 절차를 방해하지 않는다는 것을 입증하는 자료.\n7. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n8. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_11",
        "title_text": "3.2.P.3 제조\n11. 정성적 또는 정량적인 조성과 평균 질량의 변경이 없는 성상의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_11_sub_11a\") == \"변경 있음\" and step6_selections.get(\"p3_11_req_1\") == \"충족\" and step6_selections.get(\"p3_11_req_2\") == \"충족\" and step6_selections.get(\"p3_11_req_3\") == \"충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. 완제의약품의 성상.\n3. (P.3) 배치 조성에 대한 자료, 제조공정 및 공정관리의 설명자료, 해당되는 경우, 제조공정 파라미터 변경을 확인할 수 있는 자료.\n4. (P.5) 변경된 기준 및 시험방법.\n5. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_11",
        "title_text": "3.2.P.3 제조\n11. 정성적 또는 정량적인 조성과 평균 질량의 변경이 없는 성상의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_11_sub_11b\") == \"변경 있음\" and step6_selections.get(\"p3_11_req_1\") == \"충족\" and step6_selections.get(\"p3_11_req_2\") == \"충족\" and step6_selections.get(\"p3_11_req_3\") == \"충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료. \n2. 완제의약품의 성상.\n3. (P.3) 배치 조성에 대한 자료, 제조공정 및 공정관리의 설명자료, 해당되는 경우, 제조공정 파라미터 변경을 확인할 수 있는 자료.\n4. (P.5) 변경된 기준 및 시험방법.\n5. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12a\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_1\") == \"미충족\" and step6_selections.get(\"p3_12_req_3\") == \"미충족\" and step6_selections.get(\"p3_12_req_4\") == \"미충족\" and step6_selections.get(\"p3_12_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
         "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료."
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12b1\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_3\") == \"충족\" and step6_selections.get(\"p3_12_req_4\") == \"충족\" and step6_selections.get(\"p3_12_req_1\") == \"미충족\" and step6_selections.get(\"p3_12_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료.\n8. (P.8.2) 변경 후 생산규모 배치에 대한 안정성 시험 계획 및 이행서약. 단, 생물학적동등성시험을 제출하는 경우, 최소 2배치(파일럿 배치 이상)의 장기 및 가속 최소 6개월 자료. 다만, 일반제제(장용성 및 방출조절제제, 서방성제제 등 제형의 특수성이 인정되는 제제 외)에 해당하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’"
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12b2\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_3\") == \"충족\" and step6_selections.get(\"p3_12_req_4\") == \"충족\" and step6_selections.get(\"p3_12_req_1\") == \"미충족\" and step6_selections.get(\"p3_12_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료.\n5. (P.3.5) 변경 후 (위험도 평가에 따른) 실생산 규모의 3배치에 대한 공정 밸리데이션 보고서 또는 밸리데이션 실시 계획서.\n8. (P.8.2) 변경 후 생산규모 배치에 대한 안정성 시험 계획 및 이행서약. 단, 생물학적동등성시험을 제출하는 경우, 최소 2배치(파일럿 배치 이상)의 장기 및 가속 최소 6개월 자료. 다만, 일반제제(장용성 및 방출조절제제, 서방성제제 등 제형의 특수성이 인정되는 제제 외)에 해당하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’"
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12c1\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_1\") == \"충족\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_3\") == \"미충족\" and step6_selections.get(\"p3_12_req_4\") == \"미충족\" and step6_selections.get(\"p3_12_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료.\n3. (P.2) 원료의약품이 녹지 않은 상태로 존재하는 반고형 제제와 액상 제제일 때는 입도시험 또는 입자도 시험에 대한 적절한 밸리데이션 자료.\n4. (P.2 또는 R) 「의약품동등성시험기준」 [별표4] 제조소의 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n5. (P.3.5) 변경 후 (위험도 평가에 따른) 실생산 규모의 3배치에 대한 공정 밸리데이션 보고서 또는 밸리데이션 실시 계획서.\n6. (P.5.1) 변경 후 완제의약품 규격에 관한 자료.\n7. (P.5.4) 변경 전·후 최소 생산규모 1배치에 대한 배치분석 자료.\n9. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12c1\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_1\") == \"충족\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_5\") == \"충족\" and step6_selections.get(\"p3_12_req_3\") == \"미충족\" and step6_selections.get(\"p3_12_req_4\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료.\n3. (P.2) 원료의약품이 녹지 않은 상태로 존재하는 반고형 제제와 액상 제제일 때는 입도시험 또는 입자도 시험에 대한 적절한 밸리데이션 자료.\n5. (P.3.5) 변경 후 (위험도 평가에 따른) 실생산 규모의 3배치에 대한 공정 밸리데이션 보고서 또는 밸리데이션 실시 계획서.\n6. (P.5.1) 변경 후 완제의약품 규격에 관한 자료.\n7. (P.5.4) 변경 전·후 최소 생산규모 1배치에 대한 배치분석 자료.\n9. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_12",
        "title_text": "3.2.P.3 제조\n12. 완제의약품 제조공정의 일부 또는 전부에 대한 제조소 추가 또는 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_12_sub_12c2\") == \"변경 있음\" and step6_selections.get(\"p3_12_req_2\") == \"충족\" and step6_selections.get(\"p3_12_req_1\") == \"미충족\" and step6_selections.get(\"p3_12_req_3\") == \"미충족\" and step6_selections.get(\"p3_12_req_4\") == \"미충족\" and step6_selections.get(\"p3_12_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (해당하는 경우) 해당 제조소의 ‘의약품 등의 안전에 관한 규칙’ 제48조의2에 따른 제조 및 품질관리기준 적합판정서 또는 해외 제조원인 경우 제4조제1항제4호에 따른 유효기간 내의 제조증명서.\n2. (P.3.1) 수탁업소를 포함한 각 제조원, 주소, 책임소재 등의 자료.\n3. (P.2) 원료의약품이 녹지 않은 상태로 존재하는 반고형 제제와 액상 제제일 때는 입도시험 또는 입자도 시험에 대한 적절한 밸리데이션 자료.\n4. (P.2 또는 R) 「의약품동등성시험기준」 [별표4] 제조소의 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n5. (P.3.5) 변경 후 (위험도 평가에 따른) 실생산 규모의 3배치에 대한 공정 밸리데이션 보고서 또는 밸리데이션 실시 계획서.\n6. (P.5.1) 변경 후 완제의약품 규격에 관한 자료.\n7. (P.5.4) 변경 전·후 최소 생산규모 1배치에 대한 배치분석 자료.\n8. (P.8.2) 변경 후 생산규모 배치에 대한 안정성 시험 계획 및 이행서약. 단, 생물학적동등성시험을 제출하는 경우, 최소 2배치(파일럿 배치 이상)의 장기 및 가속 최소 6개월 자료. 다만, 일반제제(장용성 및 방출조절제제, 서방성제제 등 제형의 특수성이 인정되는 제제 외)에 해당하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’\n9. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_13",
        "title_text": "3.2.P.3 제조\n13. 비무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_13_sub_13a\") == \"변경 있음\" and step6_selections.get(\"p3_13_req_1\") == \"충족\" and step6_selections.get(\"p3_13_req_2\") == \"충족\" and step6_selections.get(\"p3_13_req_3\") == \"충족\" and step6_selections.get(\"p3_13_req_4\") == \"충족\" and step6_selections.get(\"p3_13_req_5\") == \"미충족\" and step6_selections.get(\"p3_13_req_6\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.3.5) (위험도 평가에 따른) 변경하고자 하는 제조 규모의 3배치에 대한 공정 밸리데이션 실시 보고서 또는 밸리데이션 실시 계획서.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치분석 자료(변경사항 13a는 연차보고시점 변경된 제조규모의 생산실적이 있는 경우 구비).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n6. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_13",
        "title_text": "3.2.P.3 제조\n13. 비무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_13_sub_13b\") == \"변경 있음\" and step6_selections.get(\"p3_13_req_1\") == \"충족\" and step6_selections.get(\"p3_13_req_2\") == \"충족\" and step6_selections.get(\"p3_13_req_3\") == \"충족\" and step6_selections.get(\"p3_13_req_4\") == \"충족\" and step6_selections.get(\"p3_13_req_5\") == \"충족\" and step6_selections.get(\"p3_13_req_6\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.3.5) (위험도 평가에 따른) 변경하고자 하는 제조 규모의 3배치에 대한 공정 밸리데이션 실시 보고서 또는 밸리데이션 실시 계획서.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치분석 자료(변경사항 13a는 연차보고시점 변경된 제조규모의 생산실적이 있는 경우 구비).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n6. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_13",
        "title_text": "3.2.P.3 제조\n13. 비무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_13_sub_13b\") == \"변경 있음\" and step6_selections.get(\"p3_13_req_1\") == \"충족\" and step6_selections.get(\"p3_13_req_2\") == \"충족\" and step6_selections.get(\"p3_13_req_3\") == \"충족\" and step6_selections.get(\"p3_13_req_4\") == \"충족\" and step6_selections.get(\"p3_13_req_6\") == \"충족\" and step6_selections.get(\"p3_13_req_5\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.3.5) (위험도 평가에 따른) 변경하고자 하는 제조 규모의 3배치에 대한 공정 밸리데이션 실시 보고서 또는 밸리데이션 실시 계획서.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치분석 자료(변경사항 13a는 연차보고시점 변경된 제조규모의 생산실적이 있는 경우 구비).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n6. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_13",
        "title_text": "3.2.P.3 제조\n13. 비무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_13_sub_13c\") == \"변경 있음\" and step6_selections.get(\"p3_13_req_1\") == \"충족\" and step6_selections.get(\"p3_13_req_2\") == \"충족\" and step6_selections.get(\"p3_13_req_3\") == \"충족\" and step6_selections.get(\"p3_13_req_4\") == \"충족\" and step6_selections.get(\"p3_13_req_5\") == \"미충족\" and step6_selections.get(\"p3_13_req_6\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.3.5) (위험도 평가에 따른) 변경하고자 하는 제조 규모의 3배치에 대한 공정 밸리데이션 실시 보고서 또는 밸리데이션 실시 계획서.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치분석 자료(변경사항 13a는 연차보고시점 변경된 제조규모의 생산실적이 있는 경우 구비).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n6. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_14",
        "title_text": "3.2.P.3 제조\n14. 무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_14_req_1\") == \"충족\" and step6_selections.get(\"p3_14_req_2\") == \"충족\" and step6_selections.get(\"p3_14_req_3\") == \"충족\" and step6_selections.get(\"p3_14_req_4\") == \"충족\" and step6_selections.get(\"p3_14_req_5\") == \"미충족\" and step6_selections.get(\"p3_14_req_6\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 현재 승인 및 신청한 배치 조성에 대한 비교표 등 변경 전·후에 관한 자료\n2. (P.3.5) (위험도 평가에 따른) 공정밸리데이션 자료 또는 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치 분석 자료(비교표 형식).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p3_14",
        "title_text": "3.2.P.3 제조\n14. 무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_14_req_1\") == \"충족\" and step6_selections.get(\"p3_14_req_2\") == \"충족\" and step6_selections.get(\"p3_14_req_3\") == \"충족\" and step6_selections.get(\"p3_14_req_4\") == \"충족\" and step6_selections.get(\"p3_14_req_5\") == \"충족\" and step6_selections.get(\"p3_14_req_6\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 현재 승인 및 신청한 배치 조성에 대한 비교표 등 변경 전·후에 관한 자료\n2. (P.3.5) (위험도 평가에 따른) 공정밸리데이션 자료 또는 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치 분석 자료(비교표 형식).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p3_14",
        "title_text": "3.2.P.3 제조\n14. 무균제제의 제조 규모 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_14_req_1\") == \"충족\" and step6_selections.get(\"p3_14_req_2\") == \"충족\" and step6_selections.get(\"p3_14_req_3\") == \"충족\" and step6_selections.get(\"p3_14_req_4\") == \"충족\" and step6_selections.get(\"p3_14_req_6\") == \"충족\" and step6_selections.get(\"p3_14_req_5\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 현재 승인 및 신청한 배치 조성에 대한 비교표 등 변경 전·후에 관한 자료\n2. (P.3.5) (위험도 평가에 따른) 공정밸리데이션 자료 또는 무균공정과 멸균 공정 밸리데이션 또는 평가결과에 관한 자료.\n3. (P.5.1) 완제의약품의 기준 및 시험방법.\n4. (P.5.4) 변경 전·후 생산 규모 완제의약품의 최소 1배치에 대한 배치 분석 자료(비교표 형식).\n5. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n6. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료."
    },
    {
        "title_key": "p3_15",
        "title_text": "3.2.P.3 제조\n15. 완제의약품의 제조공정 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_15_sub_15a\") == \"변경 있음\" and step6_selections.get(\"p3_15_req_1\") == \"충족\" and step6_selections.get(\"p3_15_req_2\") == \"충족\" and step6_selections.get(\"p3_15_req_3\") == \"충족\" and step6_selections.get(\"p3_15_req_4\") == \"충족\" and step6_selections.get(\"p3_15_req_5\") == \"충족\" and step6_selections.get(\"p3_15_req_6\") == \"충족\" and step6_selections.get(\"p3_15_req_7\") == \"충족\" and step6_selections.get(\"p3_15_req_8\") == \"미충족\" and step6_selections.get(\"p3_15_req_9\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.2) 해당되는 경우, 제조 공정의 개발에 관한 검토 자료:\n• In-Vitro 비교시험 자료, 예를 들면, 고형 제제 단위에 대한 기준 및 시험방법에서의 다시점 용출 프로파일(생산 배치 1배치, 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료. 규격을 벗어나는 경우 보고하여야 한다).\n• 원료의약품을 용해 상태 또는 비용해 상태로 함유하는 비무균 반고형 제형에 대한 In-Vitro 멤브레인 확산시험(membrane release test) 자료(생산 배치 1배치 및 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료\n• 원료의약품이 비용해 상태로 존재하는 액상 제제에 대해서는 형상의 가시적 변화를 점검하기 위한 현미경 상 자료와 입자 크기 분포의 비교 자료\n3. (P.3) 배치 조성, 제조공정과 공정관리의 설명 자료, 주요 단계와 중간체의 관리, (해당되는 경우) 공정 밸리데이션 실시 계획서 및/또는 평가에 관한 자료.\n4. (P.5) 변경 전·후 생산 규모 1배치에 대한 규격 및 시험성적서.\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n7. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_15",
        "title_text": "3.2.P.3 제조\n15. 완제의약품의 제조공정 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_15_sub_15b\") == \"변경 있음\" and step6_selections.get(\"p3_15_req_1\") == \"충족\" and step6_selections.get(\"p3_15_req_2\") == \"충족\" and step6_selections.get(\"p3_15_req_3\") == \"충족\" and step6_selections.get(\"p3_15_req_5\") == \"충족\" and step6_selections.get(\"p3_15_req_6\") == \"충족\" and step6_selections.get(\"p3_15_req_7\") == \"충족\" and step6_selections.get(\"p3_15_req_8\") == \"충족\" and step6_selections.get(\"p3_15_req_4\") == \"미충족\" and step6_selections.get(\"p3_15_req_9\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.2) 해당되는 경우, 제조 공정의 개발에 관한 검토 자료:\n• In-Vitro 비교시험 자료, 예를 들면, 고형 제제 단위에 대한 기준 및 시험방법에서의 다시점 용출 프로파일(생산 배치 1배치, 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료. 규격을 벗어나는 경우 보고하여야 한다).\n• 원료의약품을 용해 상태 또는 비용해 상태로 함유하는 비무균 반고형 제형에 대한 In-Vitro 멤브레인 확산시험(membrane release test) 자료(생산 배치 1배치 및 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료\n• 원료의약품이 비용해 상태로 존재하는 액상 제제에 대해서는 형상의 가시적 변화를 점검하기 위한 현미경 상 자료와 입자 크기 분포의 비교 자료\n3. (P.3) 배치 조성, 제조공정과 공정관리의 설명 자료, 주요 단계와 중간체의 관리, (해당되는 경우) 공정 밸리데이션 실시 계획서 및/또는 평가에 관한 자료.\n4. (P.5) 변경 전·후 생산 규모 1배치에 대한 규격 및 시험성적서.\n5. (P.8.1) 최소 2배치(파일럿배치이상)에 대한 3개월 가속 및 장기안정성 시험결과, 의약품동등성시험으로 생물학적동등성시험을 제출하는 경우 2배치(파일럿배치이상)에 대한 장기 및 가속 안정성 시험 최소 6개월 자료. 다만, 충족조건 3(일반제제에 해당한다.)를 만족하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’.\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n7. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_15",
        "title_text": "3.2.P.3 제조\n15. 완제의약품의 제조공정 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_15_sub_15b\") == \"변경 있음\" and step6_selections.get(\"p3_15_req_1\") == \"충족\" and step6_selections.get(\"p3_15_req_2\") == \"충족\" and step6_selections.get(\"p3_15_req_3\") == \"충족\" and step6_selections.get(\"p3_15_req_5\") == \"충족\" and step6_selections.get(\"p3_15_req_6\") == \"충족\" and step6_selections.get(\"p3_15_req_7\") == \"충족\" and step6_selections.get(\"p3_15_req_8\") == \"충족\" and step6_selections.get(\"p3_15_req_9\") == \"충족\" and step6_selections.get(\"p3_15_req_9\") == \"미충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.2) 해당되는 경우, 제조 공정의 개발에 관한 검토 자료:\n• In-Vitro 비교시험 자료, 예를 들면, 고형 제제 단위에 대한 기준 및 시험방법에서의 다시점 용출 프로파일(생산 배치 1배치, 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료. 규격을 벗어나는 경우 보고하여야 한다).\n• 원료의약품을 용해 상태 또는 비용해 상태로 함유하는 비무균 반고형 제형에 대한 In-Vitro 멤브레인 확산시험(membrane release test) 자료(생산 배치 1배치 및 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료\n• 원료의약품이 비용해 상태로 존재하는 액상 제제에 대해서는 형상의 가시적 변화를 점검하기 위한 현미경 상 자료와 입자 크기 분포의 비교 자료\n3. (P.3) 배치 조성, 제조공정과 공정관리의 설명 자료, 주요 단계와 중간체의 관리, (해당되는 경우) 공정 밸리데이션 실시 계획서 및/또는 평가에 관한 자료.\n4. (P.5) 변경 전·후 생산 규모 1배치에 대한 규격 및 시험성적서.\n5. (P.8.1) 최소 2배치(파일럿배치이상)에 대한 3개월 가속 및 장기안정성 시험결과, 의약품동등성시험으로 생물학적동등성시험을 제출하는 경우 2배치(파일럿배치이상)에 대한 장기 및 가속 안정성 시험 최소 6개월 자료. 다만, 충족조건 3(일반제제에 해당한다.)를 만족하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’.\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n7. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_15",
        "title_text": "3.2.P.3 제조\n15. 완제의약품의 제조공정 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_15_sub_15c\") == \"변경 있음\" and step6_selections.get(\"p3_15_req_1\") == \"미충족\" and step6_selections.get(\"p3_15_req_2\") == \"미충족\" and step6_selections.get(\"p3_15_req_3\") == \"미충족\" and step6_selections.get(\"p3_15_req_4\") == \"미충족\" and step6_selections.get(\"p3_15_req_5\") == \"미충족\" and step6_selections.get(\"p3_15_req_6\") == \"미충족\" and step6_selections.get(\"p3_15_req_7\") == \"미충족\" and step6_selections.get(\"p3_15_req_8\") == \"미충족\" and step6_selections.get(\"p3_15_req_9\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.2) 해당되는 경우, 제조 공정의 개발에 관한 검토 자료:\n• In-Vitro 비교시험 자료, 예를 들면, 고형 제제 단위에 대한 기준 및 시험방법에서의 다시점 용출 프로파일(생산 배치 1배치, 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료. 규격을 벗어나는 경우 보고하여야 한다).\n• 원료의약품을 용해 상태 또는 비용해 상태로 함유하는 비무균 반고형 제형에 대한 In-Vitro 멤브레인 확산시험(membrane release test) 자료(생산 배치 1배치 및 이전 공정의 1배치와 대조약 배치 결과에 대한 비교 자료\n• 원료의약품이 비용해 상태로 존재하는 액상 제제에 대해서는 형상의 가시적 변화를 점검하기 위한 현미경 상 자료와 입자 크기 분포의 비교 자료\n3. (P.3) 배치 조성, 제조공정과 공정관리의 설명 자료, 주요 단계와 중간체의 관리, (해당되는 경우) 공정 밸리데이션 실시 계획서 및/또는 평가에 관한 자료.\n4. (P.5) 변경 전·후 생산 규모 1배치에 대한 규격 및 시험성적서.\n5. (P.8.1) 최소 2배치(파일럿배치이상)에 대한 3개월 가속 및 장기안정성 시험결과, 의약품동등성시험으로 생물학적동등성시험을 제출하는 경우 2배치(파일럿배치이상)에 대한 장기 및 가속 안정성 시험 최소 6개월 자료. 다만, 충족조건 3(일반제제에 해당한다.)를 만족하는 경우 ‘최소 2배치(1개의 파일럿 배치 이상 포함)’.\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약.\n7. (R.1) 해당 변경 사항 외에 제조 관련 문서에 일체의 변경이 없다는 내용의 확인서(statement)."
    },
    {
        "title_key": "p3_16",
        "title_text": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_16_sub_16a\") == \"변경 있음\" and step6_selections.get(\"p3_16_req_1\") == \"충족\" and step6_selections.get(\"p3_16_req_2\") == \"충족\" and step6_selections.get(\"p3_16_req_4\") == \"충족\" and step6_selections.get(\"p3_16_req_3\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.3.3/P.3.4) 공정 중 시험의 규격 비교표 등 변경 전·후에 관한 자료."
    },
    {
        "title_key": "p3_16",
        "title_text": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_16_sub_16b\") == \"변경 있음\" and step6_selections.get(\"p3_16_req_2\") == \"충족\" and step6_selections.get(\"p3_16_req_1\") == \"미충족\" and step6_selections.get(\"p3_16_req_3\") == \"미충족\" and step6_selections.get(\"p3_16_req_4\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.2 또는 R) 「의약품동등성시험기준」 [별표3] 제조방법의 변경수준 및 제출자료의 범위에 따른 의약품동등성시험자료 및 실시한 의약품동등성시험이 동 규정에 적합함을 입증하는 자료.\n2. (P.3.3/P.3.4) 공정 중 시험의 규격 비교표 등 변경 전·후에 관한 자료.\n3. (P.3.3/P.3.4) 새로운 공정관리 시험방법을 사용하는 경우, 시험방법에 관한 자료.\n4. 새로운 시험방법을 사용하는 경우, 필요 시 밸리데이션 실시 보고서 또는 요약문.\n5. (P.5.4) 변경 전후 최소 1배치(파일럿 배치 이상)에 대한 시험성적 비교자료.\n6. 공정관리시험 및 기준의 추가, 삭제, 변경에 대한 타당성 입증 자료."
    },
    {
        "title_key": "p3_16",
        "title_text": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_16_sub_16c\") == \"변경 있음\" and step6_selections.get(\"p3_16_req_2\") == \"충족\" and step6_selections.get(\"p3_16_req_3\") == \"충족\" and step6_selections.get(\"p3_16_req_1\") == \"미충족\" and step6_selections.get(\"p3_16_req_4\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n6. 공정관리시험 및 기준의 추가, 삭제, 변경에 대한 타당성 입증 자료."
    },
    {
        "title_key": "p3_16",
        "title_text": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_16_sub_16d\") == \"변경 있음\" and step6_selections.get(\"p3_16_req_2\") == \"충족\" and step6_selections.get(\"p3_16_req_1\") == \"미충족\" and step6_selections.get(\"p3_16_req_3\") == \"미충족\" and step6_selections.get(\"p3_16_req_4\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.3.3/P.3.4) 공정 중 시험의 규격 비교표 등 변경 전·후에 관한 자료.\n3. (P.3.3/P.3.4) 새로운 공정관리 시험방법을 사용하는 경우, 시험방법에 관한 자료.\n4. 새로운 시험방법을 사용하는 경우, 필요 시 밸리데이션 실시 보고서 또는 요약문.\n5. (P.5.4) 변경 전후 최소 1배치(파일럿 배치 이상)에 대한 시험성적 비교자료.\n6. 공정관리시험 및 기준의 추가, 삭제, 변경에 대한 타당성 입증 자료."
    },
    {
        "title_key": "p3_16",
        "title_text": "3.2.P.3 제조\n16. 완제의약품 또는 반제품의 제조에 적용되는 공정관리시험 또는 공정관리시험 기준(IPC)의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p3_16_sub_16e\") == \"변경 있음\" and step6_selections.get(\"p3_16_req_2\") == \"충족\" and step6_selections.get(\"p3_16_req_1\") == \"미충족\" and step6_selections.get(\"p3_16_req_3\") == \"미충족\" and step6_selections.get(\"p3_16_req_4\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.3.3/P.3.4) 공정 중 시험의 규격 비교표 등 변경 전·후에 관한 자료.\n3. (P.3.3/P.3.4) 새로운 공정관리 시험방법을 사용하는 경우, 시험방법에 관한 자료.\n4. 새로운 시험방법을 사용하는 경우, 필요 시 밸리데이션 실시 보고서 또는 요약문.\n5. (P.5.4) 변경 전후 최소 1배치(파일럿 배치 이상)에 대한 시험성적 비교자료.\n6. 공정관리시험 및 기준의 추가, 삭제, 변경에 대한 타당성 입증 자료."
    },
    {
        "title_key": "p4_17",
        "title_text": "3.2.P.4 첨가제의 관리\n17. 첨가제 기원의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p4_17_sub_17a\") == \"변경 있음\" and step6_selections.get(\"p4_17_req_1\") == \"충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 첨가제가 식물 또는 합성 기원임을 입증하는 제조업자의 확인서(statement)."
    },
    {
        "title_key": "p4_17",
        "title_text": "3.2.P.4 첨가제의 관리\n17. 첨가제 기원의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p4_17_sub_17b\") == \"변경 있음\" and step6_selections.get(\"p4_17_req_1\") == \"미충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n2. (P.4) 첨가제의 관리에 관한 자료(기준 및 시험방법 등, 기 사용 예가 있는 경우 관련 자료 포함).\n3. (A.2) 외인성 물질에 대한 안전성 평가 자료(필요 시).\n4. 변경 전·후 첨가제의 규격(비교표 등 변경 전·후에 관한 자료) 및 성적서."
    },
    {
        "title_key": "p4_18",
        "title_text": "3.2.P.4 첨가제의 관리\n18. 별규에 해당하는 첨가제의 규격 또는 시험방법 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p4_18_req_1\") == \"충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 해당 변경에 대한 타당성 입증 자료.\n2. (P.4) 변경 전·후 규격 비교표 등 변경 전·후에 관한 자료, 변경하고자 하는 규격의 기준설정 근거자료, 시험방법에 관한 자료 및 밸리데이션 자료, 성적서."
    },
    {
        "title_key": "p4_19",
        "title_text": "3.2.P.4 첨가제의 관리\n19. 식약처장이 인정하는 공정서 규격으로 첨가제 규격의 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p4_19_req_1\") == \"충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 첨가제의 규격 비교표 등 변경 전·후에 관한 자료"
    },
    {
        "title_key": "p7_20",
        "title_text": "3.2.P.7 용기-마개 시스템\n20. 비무균제제의 직접용기 및 포장 재질, 종류 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_20_req_1\") == \"충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 용기·포장 재질 변경이 반영된 의약품의 성상.\n2. (P.2) (해당하는 경우) 현재의 포장 시스템에 비해 동등하거나 우수한 보호성을 입증하는 용기 마개 시스템의 적합성에 대한 자료. 기능성 포장을 변경하는 경우, 새로운 포장의 기능성을 입증하는 자료, 기허가 의약품에서 변경하고자 하는 용기·포장 재질의 사용례 등.\n3. (P.3) 직접 용기·포장의 재질 및 종류 변경이 반영된 제조방법 자료.\n4. (P.7) 변경하고자 하는 일차 포장 유형에 대한 정보(예 : 성상, 일차 포장 구성 성분들의 구성 재료, 규격, 성적서 등).\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p7_20",
        "title_text": "3.2.P.7 용기-마개 시스템\n20. 비무균제제의 직접용기 및 포장 재질, 종류 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_20_req_1\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 용기·포장 재질 변경이 반영된 의약품의 성상.\n2. (P.2) (해당하는 경우) 현재의 포장 시스템에 비해 동등하거나 우수한 보호성을 입증하는 용기 마개 시스템의 적합성에 대한 자료. 기능성 포장을 변경하는 경우, 새로운 포장의 기능성을 입증하는 자료, 기허가 의약품에서 변경하고자 하는 용기·포장 재질의 사용례 등.\n3. (P.3) 직접 용기·포장의 재질 및 종류 변경이 반영된 제조방법 자료.\n4. (P.7) 변경하고자 하는 일차 포장 유형에 대한 정보(예 : 성상, 일차 포장 구성 성분들의 구성 재료, 규격, 성적서 등).\n5. (P.8.1) 안정성 요약문과 결론, 2배치(파일럿배치 이상) 이상에 대한 3개월 이상의 장기보존 및 가속 시험. 해당되는 경우, 광 가혹 시험 결과.\n6. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p7_21",
        "title_text": "3.2.P.7 용기-마개 시스템\n21. 무균제제의 직접용기 및 포장 재질, 종류 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_21_req_1\") == \"충족\"\n)",
        "output_1_tag": "Cmin",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmin, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Minor change(Cmin)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 용기·포장 재질 변경이 반영된 의약품의 성상. \n2. (P.2) 현재의 포장 시스템에 비해 동등하거나 우수한 보호성을 입증하는 용기 마개 시스템의 적합성에 대한 자료. 기능성 포장을 변경하는 경우, 새로운 포장의 기능성을 입증하는 자료. \n3. (P.3) 직접 용기·포장의 재질 및 종류 변경이 반영된 제조방법 자료.\n4. (P.7) 변경하고자 하는 일차 포장 유형에 대한 정보(예 : 성상, 일차 포장 구성 성분들의 구성 재료, 규격, 성적서 등).\n5. (P.8.1) 안정성 요약문과 결론, 2배치(파일럿배치 이상)에 대한 3개월 이상의 장기보존 및 가속 시험.\n7. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p7_21",
        "title_text": "3.2.P.7 용기-마개 시스템\n21. 무균제제의 직접용기 및 포장 재질, 종류 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_21_req_1\") == \"미충족\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 용기·포장 재질 변경이 반영된 의약품의 성상. \n2. (P.2) 현재의 포장 시스템에 비해 동등하거나 우수한 보호성을 입증하는 용기 마개 시스템의 적합성에 대한 자료. 기능성 포장을 변경하는 경우, 새로운 포장의 기능성을 입증하는 자료. \n3. (P.3) 직접 용기·포장의 재질 및 종류 변경이 반영된 제조방법 자료.\n4. (P.7) 변경하고자 하는 일차 포장 유형에 대한 정보(예 : 성상, 일차 포장 구성 성분들의 구성 재료, 규격, 성적서 등).\n6. (P.8.1) 안정성 요약문과 결론, 생산 규모 3배치(최소 2배치 파일럿배치 이상)에 대한 6개월 이상의 장기보존 및 가속 시험. 해당되는 경우, 광 가혹 시험 결과.\n7. (P.8.2) 변경 후 제제의 생산규모 배치에 대한 안정성 시험 계획 및 이행서약."
    },
    {
        "title_key": "p7_22",
        "title_text": "3.2.P.7 용기-마개 시스템\n22. 직접 포장의 규격 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_22_sub_22a\") == \"변경 있음\" and step6_selections.get(\"p7_22_req_1\") == \"충족\" and step6_selections.get(\"p7_22_req_2\") == \"충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.7) 규격 비교표 등 변경 전·후에 관한 자료, 변경하고자 하는 규격의 타당성 입증 자료\n2. (P.7) 추가 또는 삭제된 시험방법에 관한 자료."
    },
    {
        "title_key": "p7_22",
        "title_text": "3.2.P.7 용기-마개 시스템\n22. 직접 포장의 규격 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_22_sub_22b\") == \"변경 있음\" and step6_selections.get(\"p7_22_req_2\") == \"충족\" and step6_selections.get(\"p7_22_req_1\") == \"미충족\"\n)",
        "output_1_tag": "AR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n \nAR, 연차보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2 제2항 및 제4항에 따른 연차보고(Annual Report, AR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (P.7) 규격 비교표 등 변경 전·후에 관한 자료, 변경하고자 하는 규격의 타당성 입증 자료\n2. (P.7) 추가 또는 삭제된 시험방법에 관한 자료."
    },
    {
        "title_key": "p7_23",
        "title_text": "3.2.P.7 용기-마개 시스템\n23. 포장단위 변경\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"p7_23_req_1\") == \"충족\" and step6_selections.get(\"p7_23_req_2\") == \"충족\"\n)",
        "output_1_tag": "IR",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nIR, 시판전보고\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2제4항 단서조항에 따른 시판전 보고(Immediate Report, IR) 수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. (R.1) 포장단위 비교표 등 변경 전·후에 관한 자료"
    },
    {
        "title_key": "ds_24",
        "title_text": "디자인스페이스(Design Space) 변경\n24. 새로운 디자인스페이스 도입 또는 허가된 디자인스페이스의 확장\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"ds_24_sub_24a\") == \"변경 있음\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경에 따른 디자인 스페이스가 타당함을 입증하는 자료(위해평가 및 다변량 연구를 포함하여 디자인 스페이스를 구성하는 다양한 파라미터들의 상호작용에 대한 연구결과를 통해 도출되었음을 입증하는 자료.\n2. 표 형식으로 디자인 스페이스를 설명한 자료. \n3. 신청 서류의 변경과 관련된 CTD 자료."
    },
    {
        "title_key": "ds_24",
        "title_text": "디자인스페이스(Design Space) 변경\n24. 새로운 디자인스페이스 도입 또는 허가된 디자인스페이스의 확장\n",
        "output_condition_all_met": "(\n    step6_selections.get(\"ds_24_sub_24b\") == \"변경 있음\"\n)",
        "output_1_tag": "Cmaj",
        "output_1_text": "📖 보고유형은 다음과 같습니다.\n\nCmaj, 변경허가(신고)\n「의약품의 품목허가‧신고‧심사 규정」 제3조의2(의약품의 허가‧신고의 변경 처리) 및 제6조(국제공통기술문서 작성)에 따라 원료의약품과 완제의약품의 제조원 또는 제조방법 중 품질에 중요한 영향을 미치는 변경허가(신고) 신청(Change, C) 대상에 해당하며, 변경사항의 중요도, 충족조건 및 제출자료 요건의 난이도 등을 고려하였을 때 Major change(Cmaj)수준의 변경사항입니다.",
        "output_2_text": "🧾 필요서류는 다음과 같습니다.\n\n1. 변경에 따른 디자인 스페이스가 타당함을 입증하는 자료(위해평가 및 다변량 연구를 포함하여 디자인 스페이스를 구성하는 다양한 파라미터들의 상호작용에 대한 연구결과를 통해 도출되었음을 입증하는 자료.\n2. 표 형식으로 디자인 스페이스를 설명한 자료. \n3. 신청 서류의 변경과 관련된 CTD 자료."
    },
]

# 각 title_key에 해당하는 STEP7_ROWS 인덱스 목록을 생성한다.
STEP7_GROUPS = {}
for idx, row in enumerate(STEP7_ROWS):
    STEP7_GROUPS.setdefault(row["title_key"], []).append(idx)

# ===== Step7 화면 =====
if st.session_state.step == 7:
    current_key = st.session_state.step6_targets[st.session_state.step7_page]
    st.markdown(
        "<h2 style='font-size:25px; text-align:center; background-color:#FAF3DB; padding:4px;'>제조방법 변경에 따른 필요서류 및 보고유형 결과</h2>",
        unsafe_allow_html=True,
    )

    # Add a blank line below the page title for better spacing
    st.write("")
    
    # title 줄바꿈 처리: \n 또는 \\n → <br> 변환
    # 내부 박스에 제목이 포함되어 있으므로 별도 표시하지 않음


    if current_key not in st.session_state.step7_results:
        st.session_state.step7_results[current_key] = []
    else:
        # Clear existing results to avoid duplicates when revisiting
        st.session_state.step7_results[current_key].clear()
        
    visible_results = []
    for idx in STEP7_GROUPS.get(current_key, []):
        row = STEP7_ROWS[idx]
        expr = row["output_condition_all_met"].lstrip("if ").rstrip(":")
        try:
            cond = eval(expr, {}, {"step6_selections": st.session_state.step6_selections})
        except Exception:
            cond = False

        if cond:
            visible_results.append(
                (row["output_1_tag"], row["output_1_text"], row["output_2_text"])
            )

    if visible_results:
        for tag, output1, output2 in visible_results:
            entry = {
                "title_text": step6_items[current_key]["title"],
                "output_1_tag": tag,
                "output_1_text": output1,
                "output_2_text": output2,
            }
            st.session_state.step7_results[current_key].append(entry)

            title = entry["title_text"].replace("\\n", "\n")
            line_parts = title.split("\n", 1)
            first_line = line_parts[0]
            second_line = line_parts[1] if len(line_parts) > 1 else ""

            html_output1 = output1.replace("\\n", "<br>").replace("\n", "<br>")
            html_output1 = html_output1.replace(
                "📖 보고유형은 다음과 같습니다.",
                "<strong>📖 보고유형은 다음과 같습니다.</strong>",
            )            
            lines = output2.split("\n")
            formatted_lines = []
            if lines:
                formatted_lines.append(lines[0].strip())
                formatted_lines.append("")
                for line in lines[1:]:
                    line = line.strip()
                    if not line:
                        continue
                    if re.match(r"^\d+\.", line):
                        formatted_lines.append(
                            f"<span style='display:block;text-indent:-1.2em;margin-left:1.2em;'>{line}</span>"
                        )
                    else:
                        formatted_lines.append(line)

            non_empty = [line for line in formatted_lines if line]

            if non_empty:
                html_output2 = non_empty[0]
                if len(non_empty) > 1:
                    html_output2 += "<br><br>" + "<br>".join(non_empty[1:])
            else:
                html_output2 = ""
            html_output2 = html_output2.replace(
                "🧾 필요서류는 다음과 같습니다.",
                "<strong>🧾 필요서류는 다음과 같습니다.</strong>",
            )

            box_html = f"""
            <div style='background-color:#fffefa;padding:10px;margin-bottom:20px;'>
                <p style='font-weight:bold;font-size:15pt;margin:0;'>{first_line}</p>
                <p style='font-weight:bold;font-size:13pt;margin:0;'>{second_line}</p>
                <br>
                <p style='font-size:11pt;margin:0;'>{html_output1}</p>
                <br>
                <br>
                <p style='font-size:11pt;margin:0;'>{html_output2}</p>
            </div>
            """
            st.markdown(box_html, unsafe_allow_html=True)

    else:
        st.write(
            "해당 변경사항에 대한 충족조건을 고려하였을 때,\n"
            "「의약품 허가 후 제조방법 변경관리 가이드라인」에서 제시하고 있는\n"
            "범위에 해당하지 않는 것으로 확인됩니다."
        )

    total_pages = len(st.session_state.step6_targets)
    st.markdown(
        f"<h6 style='text-align:center'>{st.session_state.step7_page + 1} / {total_pages}</h6>",
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)    
    with col1:
        st.button("이전단계로", on_click=go_prev_step7_page)
    with col2:
        if st.session_state.step7_page == len(st.session_state.step6_targets) - 1:
            st.button("신청양식 확인하기", on_click=go_to_step8)
        else:
            st.button("다음단계로", on_click=go_next_step7_page)

# ===== Step8: [붙임] 신청양식「의약품 허가 후 제조방법 변경관리 가이드라인(민원인 안내서)」 DOCX 생성 =====

def set_cell_font(cell, font_size=11, bold=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER    
    for paragraph in cell.paragraphs:
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.name = "Noto Serif KR"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif KR")
        paragraph.paragraph_format.line_spacing = 1.4
        paragraph.paragraph_format.space_after = Pt(0)        

def apply_document_font(doc, font_name="Noto Serif KR"):
    """Ensure every run in the document uses the specified font and cells are formatted."""

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.paragraphs:
                    cell.add_paragraph("")
                for paragraph in cell.paragraphs:
                    if not paragraph.runs:
                        paragraph.add_run("")
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_text_with_breaks(cell, text, font_size=None, bold=False, font_name="Noto Serif KR"):
    """Set cell text with line breaks preserved and optional font settings."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.4
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        if font_size is not None:
            run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)        

def clone_row(table, row_idx):
    """Clone table row at row_idx and insert below it."""
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    table._tbl.insert(row_idx + 1, new_tr)
    new_row = table.rows[row_idx + 1]
    for cell in new_row.cells:
        cell.text = ""
        set_cell_font(cell, 11)
    return new_row

def delete_row(table, row_idx):
    """Remove table row at row_idx."""
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    tbl.remove(tr)

def adjust_row_heights(table, factor=0.8):
    """Reduce table row heights by the given factor."""
    for row in table.rows:
        if row.height:
            row.height = Pt(row.height.pt * factor)

def requirement_symbol(title_key, req_key, selections):
    state = selections.get(f"{title_key}_req_{req_key}", "")
    if state == "충족":
        return "O"
    if state == "미충족":
        return "X"
    return ""
    
def create_application_docx(
    current_key,
    result,
    requirements,
    selections,
    output2_text_list,
    file_path,
    page_index=0,
    total_pages=1,
):    
    # Load template to preserve all styles and merges
    template_path = os.path.join(BASE_DIR, "제조방법변경 신청양식_empty_.docx")
    doc = Document(template_path)

    # Update the title paragraph to reflect the guideline name
    title = "[붙임] CTD 제조방법 반영 후 변경허가 신청 양식"
    p = doc.paragraphs[0]
    p.clear()
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.font.name = "Noto Serif KR"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif KR")

    table = doc.tables[0]

    # Ensure header cells use 12pt bold font
    header_cells = [
        (0, 0),
        (3, 0), (3, 1), (3, 2), (3, 3), (3, 4),
        # 4. 충족조건 헤더(행 5)
        (5, 0), (5, 4),
        # 5. 필요서류 헤더(행 11)
        (11, 0), (11, 1), (11, 2), (11, 3), (11, 4),
    ]

    req_items = sorted(requirements.items())
    max_reqs = max(1, len(req_items))
    extra_reqs = max(0, max_reqs - 4)

    doc_start = 12 + extra_reqs

    for r, c in header_cells:
        r_idx = r
        if r >= 11:
            r_idx += extra_reqs
        set_cell_font(table.cell(r_idx, c), 12, bold=True)

    # Explicitly set "5. 필요서류" header texts to ensure uniform symbols
    doc_header = table.rows[11 + extra_reqs]
    set_cell_text_with_breaks(
        doc_header.cells[0],
        "5. 필요서류 (해당하는 필요서류 기재)",
        font_size=12,
        bold=True,
    )
    set_cell_text_with_breaks(
        doc_header.cells[3],
        "구비 여부\n(O, X 선택)",
        font_size=12,
        bold=True,
    )
    doc_header.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_text_with_breaks(
        doc_header.cells[4],
        "해당 페이지 표시",
        font_size=12,
        bold=True,
    )

    # 1. 신청인: template rows 0-2, columns 2-4 hold the value area
    for r_idx, key in enumerate(["name", "site", "product"]):
        for c in range(2, 5):
            cell = table.cell(r_idx, c)
            cell.text = ""
            set_cell_font(cell, 11)

    # 2-3. 변경유형 / 신청유형 (row 4)
    change_text = result["title_text"]
    apply_text = result["output_1_tag"]

    # Due to merged cells in the template, row 4 effectively has two cells:
    # the first spans columns 0-3 and the second is column 4. Writing to each
    # index individually overwrites the merged cell contents. Write only once
    # to each unique cell to preserve the correct values.
    change_cell = table.cell(4, 0)
    set_cell_text_with_breaks(change_cell, change_text, font_size=11)
    set_cell_font(change_cell, 11)

    apply_cell = table.cell(4, 4)
    apply_cell.text = apply_text
    apply_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_font(apply_cell, 11)
        
    # 4. 충족조건 헤더 설정
    # 템플릿의 행 5를 사용하여 "4. 충족조건"과 "조건 충족 여부(O, X 선택)"를 표시
    sub_row = table.rows[5]
    sub_row.cells[0].text = "4. 충족조건"
    set_cell_font(sub_row.cells[0], 12, bold=True)
    # row 5 columns 0-3 are merged in the template. Writing to cell(5,0)
    # populates the entire merged cell. Avoid resetting cells 1-3 which would
    # overwrite the merged content and remove the text.
    set_cell_text_with_breaks(sub_row.cells[4], "조건 충족 여부\n(O, X 선택)", font_size=12, bold=True)
    set_cell_font(sub_row.cells[4], 12, bold=True)
    sub_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. 충족조건 내용 채우기 (rows 6-10 default in template)
    req_items = list(requirements.items())
    max_reqs = max(1, len(req_items))
    base_req_rows = 5
    extra_reqs = max(0, max_reqs - base_req_rows)
    for i in range(extra_reqs):
        new_row = clone_row(table, 10 + i)
        for cell in new_row.cells:
            set_cell_font(cell, 11)
    for i in range(max_reqs):
        row = 6 + i
        if i < len(req_items):
            rk, text = req_items[i]
            symbol = requirement_symbol(current_key, rk, selections)
        else:
            text = ""
            symbol = ""
        table.cell(row, 0).text = text
        set_cell_font(table.cell(row, 0), 11)
        table.cell(row, 4).text = symbol
        table.cell(row, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_font(table.cell(row, 4), 11)

    # Remove unused requirement rows if fewer than template defaults
    if max_reqs < base_req_rows:
        for i in range(base_req_rows - max_reqs):
            delete_row(table, 10 - i)

    # 5. 필요서류: rows 12-18 available in template
    doc_start = 12 + extra_reqs - max(0, base_req_rows - max_reqs)
    max_docs = max(1, len(output2_text_list))
    base_doc_rows = 7
    extra_docs = max(0, max_docs - base_doc_rows)
    # Clone additional 필요서류 행 (row 18 기준)
    for i in range(extra_docs):
        new_row = clone_row(table, doc_start + base_doc_rows - 1 + i)
        for cell in new_row.cells:
            set_cell_font(cell, 11)
    for i in range(max_docs):
        row = doc_start + i
        line = output2_text_list[i] if i < len(output2_text_list) else ""
        for c in [0, 1, 2]:
            cell = table.cell(row, c)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(line)
            set_cell_font(cell, 11)
            if re.match(r"^\d+\.", line.strip()):
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.first_line_indent = Pt(-18)
            else:
                p.paragraph_format.left_indent = None
                p.paragraph_format.first_line_indent = None
        table.cell(row, 3).text = ""
        set_cell_font(table.cell(row, 3), 11)
        table.cell(row, 4).text = ""
        set_cell_font(table.cell(row, 4), 11)

    # Remove unused 필요서류 rows
    if max_docs < base_doc_rows:
        for i in range(base_doc_rows - max_docs):
            delete_row(table, doc_start + base_doc_rows + extra_docs - 1 - i)

    adjust_row_heights(table, 0.8)

    # Add signature line and ensure page number appears after two blank lines
    sig_para = doc.add_paragraph("               책임자 성명                     서명    ")
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_para.paragraph_format.line_spacing = 1.4
    sig_para.paragraph_format.space_after = Pt(0)
    for run in sig_para.runs:
        run.font.size = Pt(14)

    # Insert two blank paragraphs before the page indicator
    for _ in range(2):
        blank_para = doc.add_paragraph("")
        blank_para.paragraph_format.space_after = Pt(0)

    # Add page indicator below the signature line
    page_para = doc.add_paragraph(f"{page_index + 1} / {total_pages}")
    page_para.paragraph_format.space_after = Pt(0)
    page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_document_font(doc, "Noto Serif KR")
    doc.save(file_path)
    return file_path

# Step 8 begins
if st.session_state.step == 8:
    step7_results = st.session_state.get("step7_results", {})
    step6_items = st.session_state.get("step6_items", {})
    step6_selections = st.session_state.get("step6_selections", {})

    # Build page list using title_key and result index
    page_list = []
    for tkey, results in step7_results.items():
        if isinstance(results, dict):
            results = [results]
            step7_results[tkey] = results
        if results:
            for idx, r in enumerate(results):
                tag = (r.get("output_1_tag") or "").strip()
                txt = (r.get("output_2_text") or "").strip()
                if tag or txt:
                    page_list.append((tkey, idx))
                else:
                    page_list.append((tkey, None))
        else:
            page_list.append((tkey, None))

    if not page_list:
        st.error("결과가 없어 Step7로 돌아갑니다.")
        st.session_state.step = 7
        render_footer()
        st.stop()

    if "step8_page" not in st.session_state:
        st.session_state.step8_page = 0

    page = st.session_state.step8_page
    total_pages = len(page_list)
    current_key, current_idx = page_list[page]
    # Render message when there is no matching result for this page
    if current_idx is None:
        st.write(
            "해당 변경사항에 대한 충족조건을 고려하였을 때,\n"
            "「의약품 허가 후 제조방법 변경관리 가이드라인」에서 제시하고 있는\n"
            "범위에 해당하지 않는 것으로 확인됩니다."
        )
    else:
        result = step7_results[current_key][current_idx]
        if not (result.get("output_1_tag") or "").strip() and not (result.get("output_2_text") or "").strip():
            st.write(
                "해당 변경사항에 대한 충족조건을 고려하였을 때,\n"
                "「의약품 허가 후 제조방법 변경관리 가이드라인」에서 제시하고 있는\n"
                "범위에 해당하지 않는 것으로 확인됩니다."
            )
        else:
            requirements = step6_items.get(current_key, {}).get("requirements", {})

            selections = {
                f"{current_key}_req_{rk}": step6_selections.get(f"{current_key}_req_{rk}", "")
                for rk in requirements
            }
            output2_text_list = [line.strip() for line in result.get("output_2_text", "").split("\n") if line.strip()]
            if output2_text_list and "필요서류" in output2_text_list[0]:
                output2_text_list = output2_text_list[1:]
            with st.spinner("⚠️ 다운로드용 파일 생성 중입니다. 잠시 기다려주세요"):
                with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    file_path = tmp.name
                    create_application_docx(
                        current_key,
                        result,
                        requirements,
                        selections,
                        output2_text_list,
                        file_path,
                    )

                with open(file_path, "rb") as f:
                    file_bytes = f.read()

            col_left, col_right = st.columns([1, 4])
            with col_left:
                st.download_button(
                    "📄 파일 다운로드",
                    file_bytes,
                    file_name=f"신청서_{current_key}_{current_idx}.docx",
                )
            with col_right:
                st.success("💡 다운로드용 파일 생성이 완료되었습니다.")
            os.remove(file_path)

            st.markdown(
                "<h5 style='text-align:left; font-size:20px;'>[붙임] CTD 제조방법 반영 후 변경허가 신청 양식</h5>",
                unsafe_allow_html=True,
            )

            # Replace newline characters in the title text for proper HTML rendering
            title_html = (
                result["title_text"].replace("\\n", "<br>").replace("\n", "<br>")
            )

            html = textwrap.dedent(

            f"""
<style>
table {{ border-collapse: collapse; width: 100%; font-family: 'Nanum Gothic', sans-serif; }}
td {{ border: 1px solid black; padding: 6px; text-align: center; vertical-align: middle; }}
.title {{ font-weight: bold; font-size: 12pt; }}
.normal {{ font-size: 11pt; }}
</style>
<table>
  <tr>
    <td class='title' rowspan='3' style='width:17.25%;background-color:#FAF3DB;'>1. 신청인</td>
    <td class='normal' style='width:27.75%'>성명</td>
    <td colspan='3' style='width:55%'></td>
  </tr>
  <tr>
    <td class='normal' style='width:27.75%'>제조소(영업소) 명칭</td>
    <td colspan='3' style='width:55%'></td>
  </tr>
  <tr>
    <td class='normal' style='width:27.75%'>변경신청 제품명</td>
    <td colspan='3' style='width:55%'></td>
  </tr>
  <tr>
    <td class='title' colspan='2' style='width:45%;background-color:#FAF3DB;'>2. 변경유형</td>
    <td class='title' colspan='3' style='width:55%;background-color:#FAF3DB;'>3. 신청 유형(AR, IR, Cmin, Cmaj 선택)</td>
  </tr>
  <tr>
    <td colspan='2' class='normal' style='width:45%'>{title_html}</td>
    <td colspan='3' class='normal' style='width:55%'>{result["output_1_tag"]}</td>
  </tr>
  <tr><td colspan='3' class='normal' style='width:60%;font-weight:bold;background-color:#FAF3DB;'>4. 충족조건</td><td colspan='2' class='normal' style='width:40%;font-weight:bold;background-color:#FAF3DB;'>조건 충족 여부(O, X 선택)</td></tr>
"""
        )

            req_items = sorted(requirements.items())
            max_reqs = max(1, len(req_items))
            for idx in range(max_reqs):
                if idx < len(req_items):
                    rk, text = req_items[idx]
                    symbol = requirement_symbol(current_key, rk, selections)
                else:
                    text = ""
                    symbol = ""
                if re.match(r"^\d+\.", text.strip()):
                    text_html = (
                        f"<span style='display:block;text-indent:-1.2em;margin-left:1.2em;'>{text}</span>"
                    )
                else:
                    text_html = text
                html += (
                    f"<tr><td colspan='3' class='normal' style='width:60%;text-align:left'>{text_html}</td>"
                    f"<td colspan='2' class='normal' style='width:40%'>{symbol}</td></tr>"
                )

            html += textwrap.dedent(
                """
  <tr>
    <td class='title' colspan='3' style='width:60%;background-color:#FAF3DB;'>5. 필요서류 (해당하는 필요서류 기재)</td>
    <td class='title' style='width:19%;background-color:#FAF3DB;'>구비 여부<br>(O, X 선택)</td>
    <td class='title' style='width:21%;background-color:#FAF3DB;'>해당 페이지 표시</td>
  </tr>
"""
            )
            max_docs = max(1, len(output2_text_list))
            for i in range(max_docs):
                line = output2_text_list[i] if i < len(output2_text_list) else ""
                if re.match(r"^\d+\.", line.strip()):
                    line_html = (
                        f"<span style='display:block;text-indent:-1.2em;margin-left:1.2em;'>{line}</span>"
                    )
                else:
                    line_html = line
                html += (
                    f"<tr>"
                    f"<td colspan='3' class='normal' style='width:60%;text-align:left'>{line_html}</td>"
                    f"<td class='normal' style='width:19%'></td>"
                    f"<td class='normal' style='width:21%'></td>"
                    f"</tr>"
                )
            html += "</table>"
            st.markdown(html, unsafe_allow_html=True)

    # Display page number and navigation for all pages
    st.markdown(
        f"<h6 style='text-align:center'>{page+1} / {total_pages}</h6>",
        unsafe_allow_html=True,
    )

    nav_left, _, nav_right = st.columns([1, 4, 1])
    with nav_left:
        if st.button("⬅ 이전"):
            if st.session_state.step8_page == 0:
                st.session_state.step = 7
                if "step8_page" in st.session_state:
                    del st.session_state["step8_page"]
            else:
                st.session_state.step8_page -= 1
    with nav_right:
        if page < total_pages - 1:
            if st.button("다음 ➡"):
                st.session_state.step8_page += 1
        else:
            if st.button("관련 자료 확인하기"):
                st.session_state.step = 9
                if hasattr(st, "rerun"):
                    st.rerun()
                else:
                    st.experimental_rerun()

if st.session_state.step == 9:
    st.markdown(
        "<h2 style='font-size:24px; text-align:center; background-color:#FAF3DB; padding:4px;'>가이드라인 및 질의응답집</h2>",
        unsafe_allow_html=True,
    )
    st.markdown("<br><br><br>", unsafe_allow_html=True)

    guideline_path = os.path.join(
        BASE_DIR, "의약품 허가 후 제조방법 변경관리 가이드라인(민원인안내서).pdf"
    )
    with open(guideline_path, "rb") as f:
        guideline_bytes = f.read()

    def mark_dl1():
        st.session_state.step9_dl1 = True

    label1 = (
        ("📂" if st.session_state.step9_dl1 else "📁")
        + " 의약품 허가 후 제조방법 변경관리 가이드라인(민원인안내서) 다운로드 하기"
    )
    st.download_button(
        label1,
        guideline_bytes,
        file_name="의약품 허가 후 제조방법 변경관리 가이드라인(민원인안내서).pdf",
        key="dl1",
        on_click=mark_dl1,
    )
    st.markdown("<br><br>", unsafe_allow_html=True)

    qna_path = os.path.join(
        BASE_DIR, "의약품 허가 후 제조방법 변경관리 질의응답집(민원인안내서).pdf"
    )
    with open(qna_path, "rb") as f:
        qna_bytes = f.read()

    def mark_dl2():
        st.session_state.step9_dl2 = True

    label2 = (
        ("📂" if st.session_state.step9_dl2 else "📁")
        + " 의약품 허가 후 제조방법 변경관리 질의응답집(민원인안내서) 다운로드 하기"
    )
    st.download_button(
        label2,
        qna_bytes,
        file_name="의약품 허가 후 제조방법 변경관리 질의응답집(민원인안내서).pdf",
        key="dl2",
        on_click=mark_dl2,
    )
    st.markdown("<br><br><br><br>", unsafe_allow_html=True)

    back_col, _ = st.columns(2)
    with back_col:
        if st.button("⬅ 이전"):
            st.session_state.step = 8

if st.session_state.step > 0:
    render_footer()
